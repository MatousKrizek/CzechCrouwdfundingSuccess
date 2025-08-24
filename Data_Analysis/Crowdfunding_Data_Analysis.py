# %% 1 Technical

targetFile = r"CrF analysis01.pickle"


# i = technical variable carrying the index free for
# the next graphical figure

# 1.1 Model pre-defintions
# List of columns (vaiables) I want to use in the analysis in the compounding version of the model
# (respecting time value of money)
f_cols_c = ["ln_c_Goal", "n_rewards", "N_videos", "c_min_contribution",
            "c_max_contribution", "c_med_contribution", "ln_descr_length",
            "Year_number", "Month"]

# List of columns (variables) I want to use in the analysis in the classical version 
# (not respecting time value of money)
f_cols_g = ["ln_Goal", "n_rewards", "N_videos", "min_contribution",
            "max_contribution", "med_contribution", "ln_descr_length",
            "Year_number", "Month"]

# 1.2 - Importing libraries
import numpy as np
from matplotlib.ticker import FixedLocator, ScalarFormatter
from sklearn.preprocessing import StandardScaler
from pygam import LogisticGAM, s, f, te, l
import re
from docx.shared import Pt
from docx import Document
import sys
from io import StringIO
from tabulate import tabulate
from sklearn.linear_model import LinearRegression
from sklearn.metrics import roc_curve, roc_auc_score, confusion_matrix
import statsmodels.formula.api as smf
import scipy.stats as stats
import seaborn as sns
import matplotlib.pyplot as plt
from datetime import datetime
import statistics
import pandas as pd
import pickle
import patsy


# 1.3 - Creating proprietary functions

def loadPcl(savedFile):
    """Loads the object pickled in the given pickle file."""
    with open(savedFile, "rb") as file:
        return pickle.load(file)

def save(object):
    """Saves the object into a predefined target Pickle file."""
    try:
        with open(targetFile, "wb") as file:
            pickle.dump(object, file, protocol=pickle.HIGHEST_PROTOCOL)
    except Exception as ex:
        print("File was not saved: ", ex)


def isolate_date(date_string):
    """Transfers date information form a HitHit string
    to datetime format."""
    x = date_string.strip()
    try:
        return (datetime.strptime(x, "%d.%m.%Y %H:%M"))
    except:
        return (pd.NaT)
    
def quantile_range(df, column, lower, upper):
    """Returns the rows of the df, where the values in a given
    column are between the given lower and upper bounds."""
    lower_quantile = df[column].quantile(lower)
    upper_quantile = df[column].quantile(upper)
    return df[(df[column] >= lower_quantile) & (df[column] <= upper_quantile)]

def success(success_rate):
    """Determines if the success rate has crossed 100 %."""
    if success_rate >= 1:
        return True
    else:
        return False

def compound(value, past_year, current_year, inflation):
    """Compounds the value up to the beginning of the inserted 
    "current year" """
    year_range = np.arange(past_year, current_year)
    cfactor_chain = inflation.loc[year_range, "cfactor"]
    total_cfactor = cfactor_chain.prod()
    compounded_value = value*total_cfactor
    return compounded_value

def num_distr_stats(df):
    """ Returns a dataframe of basic EDA statistics of all numeric
    collumns in the given dataframe df."""
    sts = {}
    for column in df.select_dtypes(include=['number']).columns:
        col_data = df[column]
        sts[column] = {
            'average': col_data.mean(),
            'median': col_data.median(),
            'average_cut': quantile_range(cf, column, 0.05, 0.95)[column].mean(),
            'standard_deviation': col_data.std(),
            'skewness': stats.skew(col_data),
            'kurtosis': stats.kurtosis(col_data)
        }
    stats_df = pd.DataFrame(sts)
    return stats_df

def summarize_numeric_variables(df, variables, round_func, output_path=None):
    """
    Summarizes numeric and boolean variables in a DataFrame.

    Parameters:
    -----------
    df : pd.DataFrame
        The input DataFrame.
    variables : list
        List of column names (numeric or boolean) to summarize.
    round_func : function
        Custom rounding function to apply to the result.
        It should take a DataFrame and return a DataFrame.
    output_path : str, optional
        File path to save the summary table as an Excel file.
        If None, the file is not saved.

    Returns:
    --------
    pd.DataFrame
        A summary table with statistics for the specified variables.
    """
    df_num = df[variables]
    summary_table = pd.DataFrame({
        'count': df_num.count(),
        'mean': df_num.mean(),
        'min': df_num.min(),
        'median': df_num.median(),
        'max': df_num.max(),
        'std': df_num.std(),
        'skewness': df_num.skew(),
        'kurtosis': df_num.kurt()
    })

    # Apply the custom rounding function
    summary_table = summary_table.applymap(round_func)

    # Optionally export to Excel
    if output_path is not None:
        summary_table.to_excel(output_path, index=True)

    return summary_table


def all_hists(df, first_figure_index):
    """Returns histogram of all numerical columns of a df."""
    i = first_figure_index
    for column in cf.select_dtypes(include=['number']).columns:
        plt.figure()
        sns.histplot(cf, x=column)
        plt.title(f'Histogram of {column}')
        plt.xlabel(column)
        plt.ylabel('Frequency')
        plt.show()
        plt.clf()
        i += 1
    return (i)


def all_boxplots(df, first_figure_index):
    """Creates box-and-whisker plots for all numeric columns
    of the given dataframe"""
    i = first_figure_index
    for col in cf.select_dtypes(include=['number']).columns:
        plt.figure(i)
        sns.boxplot(data=df[col], whis=[5, 95],
                    showmeans=True, showfliers=True, palette="Set2")
        # whiskers are now 5th and 95 th percentile
        # showfliers shows the outliers
        plt.title(f"Box-And-Whisker Plot {col}")
        plt.show()
        plt.clf()
        i += 1
    return i


def variance_inflation_factor(X, idx):
    """Returns the variance inflation factor."""
    X = np.array(X)
    mask = np.arange(X.shape[1]) != idx
    X_masked = X[:, mask]
    y = X[:, idx]
    model = LinearRegression().fit(X_masked, y)
    r_squared = model.score(X_masked, y)
    if r_squared == 1:
        return float('inf')
    return 1 / (1 - r_squared)


def calculate_vif(df):
    """Calculate VIF for each numerical column in the DataFrame."""
    # Select only numeric columns
    # .drop(columns=["ln_Collected", "ln_c_Collected"])
    numeric_df = df.select_dtypes(include=[float, int])
    # Add intercept column
    numeric_df['intercept'] = 1

    # purging -infs
    df.replace([np.inf, -np.inf], -10e+100, inplace=True)

    # Find infinite values in the DataFrame
    infinite_values = df.replace([np.inf, -np.inf], np.nan).isna()
    # Print the column and row names of infinite values
    for column in infinite_values.columns:
        for row in infinite_values.index:
            if infinite_values.at[row, column]:
                print(f"Infinite value found at column '{
                      column}', row '{row}'")

    # Calculate VIF using the inverse of the correlation matrix
    vif_data = pd.DataFrame()
    vif_data['Variable'] = numeric_df.columns
    vif_data['VIF'] = [variance_inflation_factor(
        numeric_df.values, i) for i in range(numeric_df.shape[1])]
    return vif_data


def prepare_vif_matrix(X, month_col_index, month_map, calculate_vif_func):
    # Step 1: Create a DataFrame from X
    X_df = pd.DataFrame(X, columns=[
        'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
        'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
        'ln_descr_length', 'Year_number', 'Month_code'
    ])
    
    # Step 2: Apply effect coding to the Month_code column
    month_series = X_df.iloc[:, month_col_index].astype(int).map(month_map)
    month_df = pd.get_dummies(month_series)

    # Apply sum contrast (effect coding): subtract mean across categories
    effect_coded = month_df - month_df.mean()

    # Drop the original Month_code column and replace with effect coded columns
    X_df_effect = pd.concat([
        X_df.drop(columns='Month_code'),
        effect_coded
    ], axis=1)

    # Step 3: Feed into VIF function
    vif_result = calculate_vif_func(X_df_effect)
    
    return vif_result

def custom_round(value):
    """Custom rounding function to ensure reasonable rounding of numbers for exports.
    Rounds numbers with absolute value higher than 10 to whole and others to 3 decimals."""
    if pd.isna(value):
        return value
    elif isinstance(value, (int, float)):
        if abs(value) >= 10:
            return round(value)
        else:
            return round(value, 3)
    else:
        return value


def df_to_html(df, filename='table.html', transpose=True):
    """
    Converts a dataframe to a transposed HTML table with custom rounding and saves it to a file.

        Parameters:
        df (pd.DataFrame): The dataframe to be converted to a table.
        filename (str): The name of the file to save the HTML content.
        decimals (int): The number of decimal places to round to.

        Returns:
        None
        """
    # Apply custom rounding
    rounded_df = df.applymap(custom_round)

    # Transpose the dataframe
    if transpose == True:
        transposed_df = rounded_df.transpose()

    # Convert dataframe to HTML table
    html_table = transposed_df.to_html()

    # Save to file
    with open(filename, 'w') as file:
        file.write(html_table)

    print(f"HTML table has been saved to '{filename}'.")


def summarize_dataframe(df):
    # Create an empty list to store the results
    results = []

    # Iterate through the columns of the DataFrame
    for column in df.columns:
        col_data = df[column]
        col_type = col_data.dtype

        if pd.api.types.is_numeric_dtype(col_type):
            # Calculate average and range for numerical columns
            avg = col_data.mean()
            range_lower = col_data.min()
            range_upper = col_data.max()
            mode = None  # Mode is not applicable for numerical columns
        elif pd.api.types.is_datetime64_any_dtype(col_type):
            # Calculate average and range for datetime columns
            avg = None
            range_lower = col_data.min().date()
            range_upper = col_data.max().date()
            mode = None  # Mode is not applicable for datetime columns
        else:
            # Calculate mode for categorical columns
            mode = col_data.mode()[0]
            avg = None  # Average is not applicable for categorical columns
            range_lower = None  # Range is not applicable for categorical columns
            range_upper = None  # Range is not applicable for categorical columns

        # Append the results to the list
        results.append({
            'Feature': column.capitalize(),
            'Type': col_type,
            'Average': custom_round(avg if avg is not None else None),
            'Range lower': custom_round(range_lower),
            'Range upper': custom_round(range_upper)
        })

    # Create a new DataFrame from the results list
    summary_df = pd.DataFrame(results)

    # Print the resulting DataFrame
    print(summary_df)

    # Convert the DataFrame to an HTML table
    html_table = tabulate(summary_df, headers='keys',
                          tablefmt='html', showindex=False)

    # Add some basic CSS for a clean design
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Summary Table</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: white;
                margin: 0;
                padding: 20px;
            }}
            table {{
                border-collapse: collapse;
                margin-bottom: 20px;
                table-layout: auto;
                width: auto;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
                white-space: nowrap;
            }}
            th {{
                background-color: #f2f2f2;
            }}
        </style>
    </head>
    <body>
        {html_table}
    </body>
    </html>
    """

    # Save the HTML content to a file
    with open('summary_table.html', 'w') as file:
        file.write(html_content)

    print("HTML table has been saved to 'summary_table.html'.")


def export_model_summaries_to_excel():
    with pd.ExcelWriter("Model_summaries2.xlsx", engine='openpyxl') as writer:
        # --- Logistic models ---
        for name, model in {
            "General_Logit": model_gen,
            "Compounded_Logit": model_c,
            "Interaction_Logit": model_i
        }.items():
            summary_text = model.summary().as_text()
            summary_df = model.summary2().tables[1]
            summary_df.to_excel(writer, sheet_name=name)

        # --- GAM models (capture printed summary) ---
        for name, gam_model, r2, auc in [
            ("GAM", gam, mcfadden_r2_gam, auc_gam),
            ("GAM_Smoothed", gam_s, mcfadden_r2_s, auc_s)
        ]:
            # Capture printed summary to string - deactivated - neverending
            buffer = StringIO()
            original_stdout = sys.stdout
            try:
                sys.stdout = buffer
                gam_model.summary()
            finally:
                sys.stdout = original_stdout
            summary_lines = [f"McFadden R2: {r2:.4f}", f"AUC: {
                auc:.4f}", ""] + summary_text.split("\n")
            gam_df = pd.DataFrame(summary_lines, columns=["Summary"])
            gam_df.to_excel(writer, sheet_name=name, index=False)

    print("Model summaries successfully exported to 'Model_summaries.xlsx'.")


def export_model_summary_to_word_complete(model, output_path, custom_round, effect_vars=None):
    """
    Export a statsmodels model summary to Word, including implied coefficients for effect-coded variables.

    Parameters:
    - model: fitted statsmodels Logit model
    - output_path: str (r-style path)
    - custom_round: function to round numerical values
    - effect_vars: dict — keys are prefixes (e.g. 'C(Category, Sum)[S.') and values are lists of all categories
    """
    import pandas as pd
    from docx import Document
    from docx.shared import Pt

    # Extract base summary
    summary_df = model.summary2().tables[1].copy()
    summary_df = summary_df.applymap(custom_round)

    # Recover implied coefficients
    if effect_vars:
        for prefix, categories in effect_vars.items():
            present_rows = summary_df.filter(like=prefix, axis=0)
            present_names = [i.split(prefix)[1].rstrip(']')
                             for i in present_rows.index if i.startswith(prefix)]
            missing = [cat for cat in categories if cat not in present_names]

            for miss in missing:
                # Compute implied coefficient
                present_vals = [summary_df.loc[f"{prefix}{cat}]", "Coef."] for cat in present_names if f"{
                    prefix}{cat}]" in summary_df.index]
                implied_val = -sum(present_vals)

                new_row = pd.DataFrame({
                    "Coef.": [custom_round(implied_val)],
                    "Std.Err.": ["–"],
                    "z": ["–"],
                    "P>|z|": ["–"],
                    "[0.025": ["–"],
                    "0.975]": ["–"]
                }, index=[f"{prefix}{miss}] (implied)"])

                summary_df = pd.concat([summary_df, new_row])

    summary_df = summary_df.sort_index()

    # Build Word document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading('Model Summary Table', level=2)

    table = doc.add_table(rows=1, cols=len(summary_df.columns) + 1)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    for i, col in enumerate(summary_df.columns):
        hdr_cells[i + 1].text = str(col)
        hdr_cells[i + 1].paragraphs[0].runs[0].bold = True

    # Rows
    for index, row in summary_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(index)
        for j, val in enumerate(row):
            cells[j + 1].text = str(val)

    # Save
    doc.save(output_path)
    print(f"✔️ Full model table saved to: {output_path}")


def export_model_summary_to_word_auto(model, output_path, formula, data, custom_round):
    """
    Export model summary to Word with recovered effect-coded terms based on formula and dataset.

    Parameters:
    - model: fitted statsmodels model
    - output_path: output .docx file path
    - formula: the original model formula (string)
    - data: the dataframe used to fit the model
    - custom_round: a function to round numerical outputs
    """
    summary_df = model.summary2().tables[1].copy()
    summary_df = summary_df.applymap(custom_round)

    # Step 1: Detect C(..., Sum) terms
    effect_vars = {}
    pattern = r'C\(([^,]+),\s*Sum\)'
    matches = re.findall(pattern, formula)

    for var in matches:
        categories = data[var].astype("category").cat.categories
        prefix = f'C({var}, Sum)[S.'
        effect_vars[prefix] = list(categories)

    # Step 2: Detect interactions with C(...) (only 1st-order handled)
    int_pattern = r'ln_c_Goal\s*:\s*C\(([^,]+),\s*Sum\)'
    int_matches = re.findall(int_pattern, formula)

    for var in int_matches:
        categories = data[var].astype("category").cat.categories
        prefix = f'ln_c_Goal:C({var})[T.'
        effect_vars[prefix] = list(categories)

    # Step 3: Recover implied coefficients
    for prefix, cats in effect_vars.items():
        present_rows = summary_df.filter(like=prefix, axis=0)
        present_names = [i.split(prefix)[1].rstrip(']')
                         for i in present_rows.index if i.startswith(prefix)]
        missing = [cat for cat in cats if cat not in present_names]

        for miss in missing:
            present_vals = [
                summary_df.loc[f"{prefix}{cat}]", "Coef."]
                for cat in present_names
                if f"{prefix}{cat}]" in summary_df.index
            ]
            implied_val = -sum(present_vals)

            new_row = pd.DataFrame({
                "Coef.": [custom_round(implied_val)],
                "Std.Err.": ["–"],
                "z": ["–"],
                "P>|z|": ["–"],
                "[0.025": ["–"],
                "0.975]": ["–"]
            }, index=[f"{prefix}{miss}] (implied)"])

            summary_df = pd.concat([summary_df, new_row])

    summary_df = summary_df.sort_index()

    # Step 4: Export to Word
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading('Model Summary Table', level=2)

    table = doc.add_table(rows=1, cols=len(summary_df.columns) + 1)
    table.style = 'Table Grid'

    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    for i, col in enumerate(summary_df.columns):
        hdr_cells[i + 1].text = str(col)
        hdr_cells[i + 1].paragraphs[0].runs[0].bold = True

    # Data rows
    for index, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        for j, val in enumerate(row):
            row_cells[j + 1].text = str(val)

    doc.save(output_path)
    print(f"✔️ Full summary exported to: {output_path}")


def export_model_summary_to_word_thesis(model, output_path, formula, data, custom_round):

    # Czech month names
    czech_months = [
        "Leden", "Únor", "Březen", "Duben", "Květen", "Červen",
        "Červenec", "Srpen", "Září", "Říjen", "Listopad", "Prosinec"
    ]

    # Variable name mapping
    name_map = {
        'Intercept': 'β₀ (konstanta)',
        'ln_c_Goal': 'ln(cílᵣ)',
        'n_rewards': 'nₒdₘěₙ',
        'N_videos': 'nᵥᵢdₑí',
        'c_min_contribution': 'pₘᵢₙᵣ',
        'c_max_contribution': 'pₘₐₓᵣ',
        'c_med_contribution': 'pₘₑ𝒹ᵣ',
        'ln_descr_length': 'ln(délka popisku)',
        'Year_number': 'nᵣₒₖᵤ',
        'logit_squared': 'logit²'
    }

    keys = ['Coef.', 'Std.Err.', 'z', 'P>|z|', '[0.025', '0.975]']
    column_headers = ['Proměnná', 'Odhad',
                      'Std.Chyba', 'z', 'p-hodnota', '95% CI']

    summary_df = model.summary2().tables[1].copy()
    summary_df = summary_df.applymap(custom_round)

    # Recover implied coefficients
    effect_vars = {}
    for match in re.findall(r'C\(([^,]+),\s*Sum\)', formula):
        cats = data[match].astype('category').cat.categories
        effect_vars[f'C({match}, Sum)[S.'] = list(cats)

    for match in re.findall(r'ln_c_Goal\s*:\s*C\(([^,]+),\s*Sum\)', formula):
        cats = data[match].astype('category').cat.categories
        effect_vars[f'ln_c_Goal:C({match})[T.'] = list(cats)

    for prefix, cats in effect_vars.items():
        present = summary_df.filter(like=prefix, axis=0)
        present_names = [i.split(prefix)[1].rstrip(']') for i in present.index]
        missing = [c for c in cats if c not in present_names]
        for m in missing:
            values = [summary_df.loc[f"{prefix}{c}]", "Coef."] for c in present_names if f"{
                prefix}{c}]" in summary_df.index]
            implied_val = -sum(values)
            new_row = pd.DataFrame({
                "Coef.": [custom_round(implied_val)],
                "Std.Err.": ["–"],
                "z": ["–"],
                "P>|z|": ["–"],
                "[0.025": ["–"],
                "0.975]": ["–"]
            }, index=[f"{prefix}{m}] (implied)"])
            summary_df = pd.concat([summary_df, new_row])
    summary_df = summary_df.sort_index()

    # Define sections
    section_order = [
        ('Konstanta (průsečík)', ['Intercept']),
        ('Cílová částka (hlavní efekt)', ['ln_c_Goal']),
        ('Kategorie projektu (efekty)', ['C(Category']),
        ('Interakce: kategorie × cíl', ['ln_c_Goal:C(Category']),
        # exclude intercept + goal
        ('Kontrolní proměnné', list(name_map.keys())[2:-1]),
        ('Měsíce (efekty ukončení kampaně)', ['C(Month']),
        ('Ostatní proměnné', ['logit_squared'])
    ]

    # Begin Word document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading('Shrnutí modelu – Logit', level=2)
    table = doc.add_table(rows=1, cols=len(keys) + 1)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(column_headers):
        hdr_cells[j].text = col
        hdr_cells[j].paragraphs[0].runs[0].bold = True

    used_labels = set()

    def add_section(header, patterns):
        row = table.add_row().cells
        row[0].merge(row[-1])
        row[0].text = header
        row[0].paragraphs[0].runs[0].bold = True

        for idx, row_data in summary_df.iterrows():
            idx_clean = idx.split(' (')[0]
            if idx in used_labels:
                continue

            if not any(pat in idx_clean for pat in patterns):
                continue

            # Generate Czech label
            if idx_clean in name_map:
                label = name_map[idx_clean]
            elif 'C(Category' in idx_clean:
                cat = idx_clean.split('[S.')[-1].split(']')[0]
                label = f'kategorie = {cat}'
            elif 'ln_c_Goal:C(Category' in idx_clean:
                cat = idx_clean.split('[T.')[-1].split(']')[0]
                label = f'ln(cílᵣ) × kategorie = {cat}'
            elif 'C(Month' in idx_clean:
                try:
                    m_idx = int(idx_clean.split('[S.')[-1].split(']')[0])
                    label = czech_months[m_idx - 1]
                except:
                    label = idx_clean
            else:
                continue  # skip unrecognized

            used_labels.add(idx)

            cells = table.add_row().cells
            cells[0].text = label
            for j, key in enumerate(keys):
                val = row_data.get(key, '')
                cells[j + 1].text = str(val)

    for section, patterns in section_order:
        add_section(section, patterns)

    doc.add_paragraph(
        "* Koeficient označený jako „implied“ byl dopočítán na základě efektového kódování (součet koeficientů = 0).")
    doc.save(output_path)

    return f"Dokument byl uložen do: {output_path}"


def export_model_summary_to_word_thesis_m(model, output_path, formula, data, custom_round=custom_round, model_name="Model_(X)"):

    # Czech month names
    czech_months = [
        "Leden", "Únor", "Březen", "Duben", "Květen", "Červen",
        "Červenec", "Srpen", "Září", "Říjen", "Listopad", "Prosinec"
    ]

    # Variable name mapping
    name_map = {
        'Intercept': 'β₀ (konstanta)',
        'ln_c_Goal': 'ln(cílᵣ)',
        'ln_Goal': 'ln(cíl)',
        'n_rewards': 'n_odměn',
        'N_videos': 'n_videí',
        'c_min_contribution': 'p_min_r',
        'c_max_contribution': 'p_max_r',
        'c_med_contribution': 'p_med_r',
        'min_contribution': 'p_min',
        'max_contribution': 'p_max',
        'med_contribution': 'p_med',
        'ln_descr_length': 'ln(délka popisku)',
        'Year_number': 'nᵣₒₖᵤ',
        'logit_squared': 'logit²'
    }

    keys = ['Coef.', 'Std.Err.', 'z', 'P>|z|', '[0.025', '0.975]']
    column_headers = ['Proměnná', 'Odhad', 'Std.Chyba',
                      'z', 'p-hodnota', '[0.025', '0.975]']

    summary_df = model.summary2().tables[1].copy()
    summary_df = summary_df.applymap(custom_round)

    # Recover implied coefficients
    effect_vars = {}
    # non-interaction categories formula
    for match in re.findall(r'C\(([^,]+),\s*Sum\)', formula):
        cats = data[match].astype('category').cat.categories
        effect_vars[f'C({match}, Sum)[S.'] = list(cats)

    # interaction vars formula
    for match in re.findall(r'ln_c_Goal\s*:\s*C\(([^,]+),\s*Sum\)', formula):
        cats = data[match].astype('category').cat.categories
        effect_vars[f'ln_c_Goal:C({match})[T.'] = list(cats)

    for prefix, cats in effect_vars.items():
        present = summary_df.filter(like=prefix, axis=0)
        present_names = [i.split(prefix)[1].rstrip(']') for i in present.index]
        # finds missing categories
        missing = [c for c in cats if str(c) not in present_names]
        for m in missing:
            values = [summary_df.loc[f"{prefix}{c}]", "Coef."] for c in present_names if f"{
                prefix}{c}]" in summary_df.index]
            implied_val = -sum(values)
            new_row = pd.DataFrame({
                "Coef.": [custom_round(implied_val)],
                "Std.Err.": ["–"],
                "z": ["–"],
                "P>|z|": ["–"],
                "[0.025": ["–"],
                "0.975]": ["–"]
            }, index=[f"{prefix}{m}] (implied)"])
            summary_df = pd.concat([summary_df, new_row])

    # summary_df = summary_df.sort_index()
    print(summary_df)

    # Define sections
    section_order = [
        ('Konstanta (průsečík)', ['Intercept']),
        ('Cílová částka (hlavní efekt)', ['ln_c_Goal', "ln_Goal"]),
        ('Interakce: kategorie × cíl', ['ln_c_Goal:C(Category']),
        ('Kategorie projektu (efekty)', ['C(Category']),
        # exclude intercept + goal
        ('Kontrolní proměnné', list(name_map.keys())[2:-1]),
        ('Měsíce (efekty ukončení kampaně)', ['C(Month']),
        ('Ostatní proměnné', ['logit_squared'])
    ]

    # Begin Word document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading(f'Shrnutí modelu – {model_name}', level=2)
    table = doc.add_table(rows=1, cols=len(keys) + 1)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(column_headers):
        hdr_cells[j].text = col
        hdr_cells[j].paragraphs[0].runs[0].bold = True

    used_labels = set()

    def add_section(header, patterns):
        row = table.add_row().cells
        row[0].merge(row[-1])
        row[0].text = header
        row[0].paragraphs[0].runs[0].bold = True

        for idx, row_data in summary_df.iterrows():
            idx_clean = idx.split(' (')[0]
            if idx in used_labels:
                continue        # escapes for loop in case of repeating

            if not any(pat in idx_clean for pat in patterns):
                continue        # "patterns" = English names of variables
                # escapes loop if patterns are already used up
            # Generate Czech label
            if idx_clean in name_map:
                label = name_map[idx_clean]
            elif 'ln_c_Goal:C(Category' in patterns and 'ln_c_Goal:C(Category' in idx_clean:
                cat = idx_clean.split('[S.')[-1].split(']')[0]
                label = f'ln(cílᵣ) × kategorie = {cat}'
            elif "C(Category" in patterns and 'C(Category' in idx_clean:
                cat = idx_clean.split('[S.')[-1].split(']')[0]
                label = f'kategorie = {cat}'

            elif 'C(Month' in idx_clean:
                try:
                    m_idx = int(idx_clean.split('[S.')[-1].split(']')[0])
                    label = czech_months[m_idx - 1]
                except:
                    label = idx_clean
            else:
                continue  # skip unrecognized

            used_labels.add(idx)

            cells = table.add_row().cells
            cells[0].text = label
            for j, key in enumerate(keys):
                val = row_data.get(key, '')
                cells[j + 1].text = str(val)

    for section, patterns in section_order:
        add_section(section, patterns)

    doc.add_paragraph(
        "* Koeficient označený jako „implied“ byl dopočítán na základě efektového kódování (součet koeficientů = 0).")
    doc.save(output_path)

    return f"Dokument byl uložen do: {output_path}"

# %% 2 Data Collection

# 2 Data Collection


# 2.1 - Extracting data from HitHit - currently using the improved data cleaner
data_HH = loadPcl(r"HitHitDF.pickle")

print(data_HH["Mode"])

print(data_HH["End_date"])

data_HH["empirical_success_rate"] = data_HH["Amount collected"]/data_HH["Goal"]
data_HH["success"] = data_HH["empirical_success_rate"].apply(
    lambda x: success(x))
data_HH["n_rewards"] = data_HH["Reward_prices"].apply(lambda x: len(x))
data_HH["video"] = data_HH["N_videos"].apply(
    lambda x: True if x > 0 else False)
data_HH["min_contribution"] = data_HH["Reward_prices"].apply(lambda x: min(x))
data_HH["max_contribution"] = data_HH["Reward_prices"].apply(lambda x: max(x))
data_HH["med_contribution"] = data_HH["Reward_prices"].apply(
    lambda x: statistics.median(x))
data_HH["descr_lenght"] = data_HH["Description"].apply(lambda x: len(x))
data_HH["End_date_DTT"] = data_HH["End_date"].apply(lambda x: isolate_date(x))

# 2.2 - Extracting data from Startovač?


# 2.3 - Creating the dataset

# the following dataset should be minimalistic
analysis_collumns = ["Title", "success", "Amount collected", "Goal", "Categories", "n_rewards", "video", "N_videos", "min_contribution", "max_contribution",
                     "med_contribution", "descr_lenght", "End_date_DTT"]

cf = data_HH[analysis_collumns].rename(columns={"Amount collected": "Collected",
                                                "End_date_DTT": "End_date", "descr_lenght": "descr_length"})


# 2.4 - Additional data
# inflation = pd.Series([1.5,1.9,], index=np.arange(2010,2024,1))
inflation = pd.read_csv("inflace.csv")
inflation = inflation.set_index("label")
inflation.columns = ["inflation"]
inflation["cfactor"] = inflation["inflation"]/100+1

# %% 3 Initial data cleaning
# 3 Initial data cleaning

summary_cf = cf.describe()
print(summary_cf)
df_to_html(summary_cf)

# Convert dataframe to HTML table
html_table = summary_cf.to_html()

# Save to file
with open('summary_cf.html', 'w') as file:
    file.write(html_table)

# Pre-NA filling calculations
# 3.2 Adding additional collumns
cf["Year"] = cf["End_date"].dt.year
cf["Year_number"] = 2024 - cf["Year"]  # years before 2024
cf["Month"] = cf["End_date"].dt.month
cf['Month'] = cf['Month'].astype('category')
cf["S-rate"] = cf["Collected"]/cf["Goal"]

summary_cf2 = cf.describe()


# 3.1 Handling NAs
cf['End_date'].fillna(method='ffill', inplace=True)
print(cf["End_date"])

# 3.2 Adding additional collumns
cf["Year"] = cf["End_date"].dt.year
cf["Year_number"] = 2024 - cf["Year"]  # years before 2024
cf["Month"] = cf["End_date"].dt.month
cf['Month'] = cf['Month'].astype('category')
cf["S-rate"] = cf["Collected"]/cf["Goal"]

# Compounding values
current_year = 2025
monetary = ['Collected', "Goal", "min_contribution",
            "max_contribution", "med_contribution"]
monetary_extended = monetary[:]
for col in monetary:
    cf[f"c_{col}"] = cf.apply(lambda row:
                              compound(row[col], row['Year'], current_year, inflation), axis=1)
    monetary_extended.append(f"c_{col}")

# Adding logarithms
numeric_columns = monetary_extended + ["descr_length", "n_rewards"]
for numeric_column in numeric_columns:
    cf[f"ln_{numeric_column}"] = np.log(cf[numeric_column])


# 3.3 Pre-EDA encoding
cf['success'] = cf['success'].astype(int)
cf['video'] = cf['video'].astype(int)

# 3.4 Simplifiing categories
# kategory "antivir" should be isolated due to its special circumstances
antivir = ["Antivir"]
# kategory "energy" is not relevant, it is only an add-on as a part of a special action
media_art = ["Hudba", "Média", "Film", "Umění",
             "Literatura", "Tanec", "Divadlo", "Fotografie"]

# categories "komunita", "nadace vodafone" and "impact hub" hide sustainable projects
altruistic = ["Komunita", "Impact Hub", "Nadace Vodafone"]

# category "design", "jídlo" (cukrárny, kavárny...), "hry" and "technologie" typicially include for-profit ideas
gastro = ["Jídlo"]
technology = ["Technologie"]
games = ["Hry"]
design = ["Design", "Móda"]  # clothing = products as well

# specific categories are education and sport
edu = ["Vzdělávání"]
sport = ["Sport"]


# Decision tree:
# 1. Filter out "Antivir" - isolate all covid specialities
# 2. Filter out "Games"     - games developers tend to put "technology" as well - not the other way around
# 3. Filter out "Technology"
# 4. Filter out "Design" - those are basicly non-tech products
# At this point, major for-profit categories should be out
# 5. Filter out "Media-art"
# 6. Filter out "Gastro" after all books about food are out
# 7. Filter out "Sport"
# At this point, semi-for-profit categories should be out
# 8. Filter out "Altruistic" - finally possible without catching on greenwashing
# 9. Filter out "Education" - those who remain here have education as their primary category
# gastro, edu and sport should remain - lets see how many that is
# 10. Empirically tested - the two "Ecology only", who remain here, are in fact products


def check_belonging(x, category):
    common_values = set(category).intersection(x)
    if common_values:
        return True
    else:
        return False


def categorize(x):                  
    if check_belonging(x, antivir):
        return "Antivirus"
    elif check_belonging(x, games):
        return "Hry"
    elif check_belonging(x, technology):
        return "Technologie"
    elif check_belonging(x, design):
        return "Designové a ostatní produkty"
    elif check_belonging(x, media_art):
        return "Umění a multimédia"
    elif check_belonging(x, gastro):
        return "Gastronomie"
    elif check_belonging(x, sport):
        return "Sport"
    elif check_belonging(x, altruistic):
        return "Komunitní a neziskové"
    elif check_belonging(x, edu):
        return "Vzdělávací"
    else:
        # Projects which remained in this group have been checked manually
        return "Designové a ostatní produkty"


cf["Category"] = cf["Categories"].apply(categorize).astype("category")
cat_frequency = cf['Category'].value_counts()


# now change "Year" to return years from the beginning and not the end
max_year = cf["Year_number"].max()
cf["Year_number"] = cf['Year_number'].apply(lambda x: max_year - x + 1)

# %% 4 Exploratory data analysis (EDA)
# 4 Exploratory data analysis (EDA)

# 4.1 Eploration through statistics
pre_treatment_stats = num_distr_stats(cf)

# vector of variances/covariation matrix
covariance_matrix = cf.select_dtypes(include=['number']).cov()

# correlation matrix
correlation_matrix = cf.select_dtypes(include=['number']).corr()
# it seems there will be no major colinearity problems

vif_data_g = calculate_vif(cf[f_cols_g])
vif_data_c = calculate_vif(cf[f_cols_c])


# %% 4.2 Graphical exploration
# 4.2 Graphical exploration
i = 1

# all histograms (used to gain direction)
# i = all_hists(cf, first_figure_index = i)

# histogram of distriubtion of success among projects
plt.figure(i, figsize=(12, 6))
sns.histplot(cf, x="S-rate")
plt.ylabel("Četnost")
plt.xlabel("Vybraný podíl cílové částky")
# Get current x-axis ticks and add 100% if it's not already there
current_ticks = plt.gca().get_xticks()
current_ticks = list(current_ticks) + [1.0]
current_ticks = [x for x in current_ticks if x >= 0]
plt.xticks(current_ticks, [f'{tick:.0%}' for tick in current_ticks])
plt.xticks(rotation=45)
plt.show()
plt.clf()
i += 1

# histogram of goal distribution
plt.figure(i)
sns.histplot(cf, x="Goal")
plt.show()
plt.clf()
i += 1

# histogram of the distribution of success rates
plt.figure(i)
sns.histplot(data=cf, x="S-rate", multiple="stack")  # hue = "Categories",
plt.show()
plt.clf()
i += 1


# boxplots for all variables
plt.figure(i, figsize=(30, 20))
numeric_columns = cf.select_dtypes(include=['number']).columns.tolist()
cf.boxplot(column=numeric_columns)
plt.xticks(rotation=45)
plt.title('Box-and-Whisker Plot for All Columns')
plt.show()
plt.clf()
i += 1

# boxplots per variable
i = all_boxplots(cf, first_figure_index=i)


# scatter plot of distribution of projects in time
cf["success_binary"] = cf["success"].apply(lambda x: bool(x))
plt.figure(i)               # Create a scatter plot with customized legend
sns.relplot(cf, x="End_date", y="Goal", hue="success_binary",
            hue_order=[True, False])
# legend = False)
plt.legend(title='', labels=['Úspěch', 'Neúspěch'])  # Customize the legend
plt.xlabel("Cílová částka")
plt.ylabel("Skončení projektu")
plt.show()
plt.clf()
i += 1

max_demander = cf.loc[cf["Goal"].idxmax()]["Title"]
print(max_demander)  # je to e-bioneta

# scatter plot of distribution of projects in time by category
plt.figure(i)
sns.relplot(cf, x="End_date", y="Goal", hue="Category")
plt.show()
plt.clf()
i += 1

# goal of individual categories (box plot)
plt.figure(i)
sns.boxplot(data=cf, x="Category", y="Goal", hue="success")
plt.show()
plt.clf()
i += 1

# histogram of project in time and their success rate
plt.figure(i)
sns.histplot(data=cf, x="Year", hue="success", binwidth=1,
             kde=False, element="bars", multiple="stack")
# improve the location of the year under the plot
plt.xticks(np.arange(cf['Year'].min() + 0.5, cf['Year'].max() + 1.5, 1),
           labels=np.arange(cf['Year'].min(), cf['Year'].max() + 1).astype(int))
plt.xlabel('Year')
plt.ylabel('Number of Campaigns')
plt.title('Campaigns per Year')
plt.show()  # za 2024 máme jen velmi málo pozorování
plt.clf()
i += 1

# plot of the ratio of succesful projects in years
success_ratio = cf.groupby("Year")["success"].mean().reset_index()
plt.figure(i)
sns.scatterplot(data=success_ratio, x='Year', y='success')
plt.xlabel('Year')
plt.ylabel('Success Ratio')
plt.title('Success Ratio of Projects by Year')
plt.show()
plt.clf()
i += 1

# histogram of project in time showing impact of covid
plt.figure(i)
cf_antivir = cf.copy()
cf_antivir["Category Antivir"] = cf["Categories"].apply(
    lambda x: True if "Antivir" in x else False)
cf_antivir = cf_antivir.sort_values(by="Category Antivir", ascending=False)

sns.histplot(data=cf_antivir, x="Year", hue="Category Antivir", binwidth=1,
             kde=False, element="bars", multiple="stack", hue_order=[True, False])
plt.xticks(np.arange(cf['Year'].min() + 0.5, cf['Year'].max() + 1.5, 1),
           labels=np.arange(cf['Year'].min(), cf['Year'].max() + 1).astype(int))
plt.xlabel('Year')
plt.ylabel('Number of Campaigns')
plt.title('Campaigns per Year')
plt.show()
plt.clf()
i += 1

# graph of ratio of succesful projects in years excluding "Antivir" cat
success_ratio = cf[cf["Category"] != "Antivir"].groupby(
    "Year")["success"].mean().reset_index()
plt.figure(i)
sns.scatterplot(data=success_ratio, x='Year', y='success')
plt.xlabel('Year')
plt.ylabel('Success Ratio')
plt.title('Success Ratio of Projects by Year excl. Antivir')
plt.show()
plt.clf()
i += 1


# plot of ratio of succesful projects for all categories in years
success_rate = cf.groupby(['Year', 'Category'])['success'].mean().reset_index()
plt.figure(i, figsize=(10, 6))
sns.lineplot(data=success_rate, x='Year', y='success',
             hue='Category', marker="o", linewidth=3)
plt.xlabel('Year')
plt.ylabel('Success Rate')
plt.title('Success Rate of All Categories by Year')
plt.legend(title='Category')
plt.show()
plt.clf()
i += 1

# another attempt with a heatmap
plt.figure(i, figsize=(12, 8))
pivot_table = success_rate.pivot(
    index='Category', columns='Year', values='success')
sns.heatmap(pivot_table, annot=True, cmap='YlGnBu', cbar=True)
plt.xlabel('Year')
plt.ylabel('Category')
plt.title('Success Rate of All Categories by Year')
plt.show()
plt.clf()
i += 1

# graph of monthly ratio of succesful projects
monthly_success_ratio = cf.groupby(['Year', 'Month'])[
    'success'].mean().reset_index()
plt.figure(i, figsize=(12, 8))
sns.lineplot(data=monthly_success_ratio, x='Month',
             y='success', hue='Year', marker='o', linewidth=2.5)
plt.xlabel('Month')
plt.ylabel('Success Ratio')
plt.title('Monthly Success Ratio of Projects by Year')
plt.legend(title='Year')
plt.show()
plt.clf()
i += 1

# %% 5 Post-EDA adjustments
# 5 Post-EDA adjustments
# New data frame cft

# 5.1 - Additional categories coding


# 5.2 - Outliers treatment

# Calculate Z-scores
z_scores = np.abs(stats.zscore(cf.select_dtypes(include=[np.number])))

# Define a threshold for identifying outliers
threshold = 3

# Identify rows with outliers
outliers = (z_scores > threshold).any(axis=1)

# Remove outliers
cf_cleaned = cf[~outliers]

print(f"Number of outliers removed: {outliers.sum()}")
print(f"Shape of cleaned DataFrame: {cf_cleaned.shape}")


# %% 6 Pre-modelling check
# 6 Pre-modelling check
num_vars2 = ["ln_Goal", "n_rewards", "N_videos", "min_contribution",
             "max_contribution", "med_contribution", "ln_descr_length", "Year_number"]

# 6.1 statistics

# description + stats
post_treatment_stats_2 = summarize_numeric_variables(
    cf_cleaned, num_vars2, custom_round, "table23.xlsx")


# vector of variances/covariation matrix
covariance_matrix_2 = cf_cleaned.select_dtypes(include=['number']).cov()

# correlation matrix
correlation_matrix_2 = cf_cleaned.select_dtypes(include=['number']).corr()
# it seems there will be no major colinearity problems

# 6.2 histograms
i = all_hists(cf_cleaned, first_figure_index = i)

# 6.3 boxplots
i = all_boxplots(cf_cleaned, first_figure_index = i)

# 6.4 Categories distribution
plt.figure(i, figsize=(12, 6))
ax = sns.countplot(data=cf_cleaned, x="Category")
plt.xticks(rotation=45, ha="right")
plt.xlabel('Kategorie')
plt.ylabel('Četnost projektů')
# Add the number of observations on top of the bars
for p in ax.patches:
    rounded_height = round(p.get_height())
    ax.annotate(f'{rounded_height}', (p.get_x() + p.get_width() / 2., p.get_height()),
                ha='center', va='center', xytext=(0, 10), textcoords='offset points')
plt.show()
plt.clf()
i += 1

# 6.5 Categories distribution
plt.figure(i, figsize=(16, 8))
palette = sns.color_palette()
ax = sns.countplot(data=cf_cleaned, x="Month", color=palette[1])
plt.xticks(rotation=45, ha="right")
plt.xlabel('Měsíc')
plt.ylabel('Četnost projektů')
# Add the number of observations on top of the bars
for p in ax.patches:
    rounded_height = round(p.get_height())
    ax.annotate(f'{rounded_height}', (p.get_x() + p.get_width() / 2., p.get_height()),
                ha='center', va='center', xytext=(0, 10), textcoords='offset points')
plt.show()
plt.clf()
i += 1

# %% 7 Additional graphics for export
# 7 Additional graphics for export


# 7.1 goal of individual categories (box plot)
plt.figure(i, figsize=(12, 6))
sns.boxplot(data=cf_cleaned, x="Category", y="Goal")
plt.xticks(rotation=45, ha="right")
plt.xlabel('Category')
plt.ylabel('Goal')              # G1.1a
# plt.tight_layout()
plt.show()
plt.clf()
i += 1

# 7.2 goal of individual categories (box plot) - real value 
plt.figure(i, figsize=(12, 6))
sns.boxplot(data=cf_cleaned, x="Category", y="c_Goal")
plt.xticks(rotation=45, ha="right")
plt.xlabel('Kategorie')
plt.ylabel('Cílová částka (v reálné hodnotě k roku 2024)')              # G1.1a
# plt.tight_layout()
plt.show()
plt.clf()
i += 1

# 7.3 goal of individual categories by success (box plot, real value) - obr. 12
# Define custom labels and palette
hue_labels = {0: "Neúspěšné", 1: "Úspěšné"}
custom_palette = {0: sns.color_palette()[1], 1: sns.color_palette()[
    0]}  # orange and blue
plt.figure(figsize=(12, 6))
# Boxplot grouped by success status (hue) within each category
sns.boxplot(
    data=cf_cleaned,
    x="Category", y="c_Goal",
    hue="success",
    palette=custom_palette
)
# Replace legend labels
handles, labels = plt.gca().get_legend_handles_labels()
new_labels = [hue_labels[int(lbl)] for lbl in labels]
plt.legend(handles, new_labels, title="Výsledek")
# Axis formatting
plt.xticks(rotation=45, ha="right")
plt.xlabel('Kategorie')
plt.ylabel('Cílová částka (v reálné hodnotě k roku 2024)')
# plt.tight_layout()
plt.show()
plt.clf()


# 7.4 Line plot of share of succcesful project per goal quantile
# Calculate quantiles
cf_cleaned['ln_c_Goal_quantile'] = pd.qcut(cf_cleaned['ln_c_Goal'],
                                           q=10, labels=False)
# Group by quantiles and calculate average success rate
quantile_success = cf_cleaned.groupby('ln_c_Goal_quantile'
                                      )['success'].mean().reset_index()
# Extract quantile intervals
quantile_intervals = sorted(pd.qcut(cf_cleaned['ln_c_Goal'],
                                    q=10, duplicates='drop').unique())
quantile_intervals_rounded = [pd.Interval(round(interval.left),
                                          round(interval.right)) for interval in quantile_intervals]
# Plot the data
plt.figure(i, figsize=(10, 6))
sns.lineplot(data=quantile_success, x='ln_c_Goal_quantile',
             y='success', marker='o', linewidth=3)
plt.xlabel('Decil požadované reálné cílové částky')
plt.ylabel('Podíl úspěšných projektů')           # G0.9
plt.title('Podíl úspěšných projektů v jednotlivých kvantilech cílové částky')

# Set xticks to show all quantiles
plt.xticks(ticks=quantile_success['ln_c_Goal_quantile'],
           labels=[str(interval) for interval in quantile_intervals_rounded],
           rotation=45)
# add quantile intervals
# for i, interval in enumerate(quantile_intervals):
#   plt.text(i, quantile_success['success'][i], str(interval), horizontalalignment='center', verticalalignment='bottom')
plt.show()
plt.clf()
i += 1


# 7.5 Scatter plot of share of succcesful project per goal quantile
# Calculate quantiles
cf_cleaned['c_Goal_quantile'] = pd.qcut(cf_cleaned['c_Goal'],
                                        q=10, labels=False)
# Group by quantiles and calculate average success rate
quantile_success = cf_cleaned.groupby('c_Goal_quantile'
                                      )['success'].mean().reset_index()
# Extract quantile intervals
quantile_intervals = sorted(pd.qcut(cf_cleaned['c_Goal'],
                                    q=10, duplicates='drop').unique())
quantile_intervals_rounded = [pd.Interval(round(interval.left),
                                       round(interval.right)) for interval in quantile_intervals]
# Plot the data
plt.figure(i, figsize=(10, 6))
sns.lineplot(data=quantile_success, x='c_Goal_quantile',
             y='success', marker='o', linewidth=0, markersize=10)
plt.xlabel('Decil požadované reálné cílové částky')
plt.ylabel('Podíl úspěšných projektů')           # G0.9
plt.title('Podíl úspěšných projektů v jednotlivých kvantilech cílové částky')
# Set xticks to show all quantiles
plt.xticks(ticks=quantile_success['c_Goal_quantile'],
           labels=[str(interval) for interval in quantile_intervals_rounded],
           rotation=45)
# add quantile intervals
# for i, interval in enumerate(quantile_intervals):
#   plt.text(i, quantile_success['success'][i], str(interval), horizontalalignment='center', verticalalignment='bottom')
plt.show()
plt.clf()
i += 1


# Define color
blue_color = sns.color_palette("deep")[0]


# 7.6 Line plot of share of succcesful project per goal quantile plus quantile size
cf_cleaned['c_Goal_quantile'] = pd.qcut(
    cf_cleaned['c_Goal'], q=10, labels=False)
quantile_success = cf_cleaned.groupby('c_Goal_quantile')[
    'success'].mean().reset_index()
quantile_intervals = pd.qcut(
    cf_cleaned['c_Goal'], q=10, duplicates='drop').cat.categories
quantile_labels = [f"{int(round(i.left))}–{int(
    round(i.right))}" for i in quantile_intervals]
quantile_counts = cf_cleaned['c_Goal_quantile'].value_counts().sort_index()
fig, ax1 = plt.subplots(figsize=(10, 6))
ax1.plot(quantile_success['c_Goal_quantile'], quantile_success['success'],
         color=blue_color, marker='o', linewidth=1, markersize=8)
ax1.set_xlabel('Decil požadované reálné cílové částky')
ax1.set_ylabel('Podíl úspěšných projektů', color=blue_color)
ax1.tick_params(axis='y', labelcolor=blue_color)
ax1.set_xticks(quantile_success['c_Goal_quantile'])
ax1.set_xticklabels(quantile_labels, rotation=45)
ax2 = ax1.twinx()
ax2.bar(quantile_success['c_Goal_quantile'], quantile_counts,
        width=0.5, color='lightgray', alpha=0.6)
ax2.set_ylabel("Počet projektů", color='gray')
ax2.tick_params(axis='y', labelcolor='gray')
fig.legend(loc="upper right")
plt.title("Úspěšnost podle decilů cílové částky")
plt.tight_layout()
plt.show()

# 7.7 Scatter plot of share of succcesful project per goal quantile plus quantile size
# Create quantiles and calculate statistics
cf_cleaned['c_Goal_quantile'] = pd.qcut(
    cf_cleaned['c_Goal'], q=10, labels=False)
quantile_success = cf_cleaned.groupby('c_Goal_quantile')[
    'success'].mean().reset_index()
quantile_intervals = pd.qcut(
    cf_cleaned['c_Goal'], q=10, duplicates='drop').cat.categories
quantile_labels = [f"{int(round(i.left))}–{int(
    round(i.right))}" for i in quantile_intervals]
quantile_counts = cf_cleaned['c_Goal_quantile'].value_counts().sort_index()
# Plotting
fig, ax1 = plt.subplots(figsize=(10, 6))
# Secondary axis for histogram
ax2 = ax1.twinx()
ax2.bar(quantile_success['c_Goal_quantile'], quantile_counts,
      width=0.9, color='lightgray', alpha=0.6, label='Počet projektů', zorder=1)
ax2.set_ylabel("Počet projektů", color='gray')
ax2.tick_params(axis='y', labelcolor='gray')
# Bring ax1 to front
ax1.set_zorder(ax2.get_zorder() + 1)
ax1.patch.set_visible(False)  # Hide the background patch of ax1
# Scatterplot on primary axis
ax1.scatter(quantile_success['c_Goal_quantile'], quantile_success['success'],
            color=blue_color, s=50, label='Úspěšnost', zorder=2)
ax1.set_xlabel('Decil požadované reálné cílové částky')
# , fontweight='bold')
ax1.set_ylabel('Podíl úspěšných projektů', color="black")
ax1.tick_params(axis='y', labelcolor="black")
ax1.set_xticks(quantile_success['c_Goal_quantile'])
ax1.set_xticklabels(quantile_labels, rotation=45)
fig.legend(loc="upper right")
plt.title("Úspěšnost podle decilů cílové částky")
plt.tight_layout()
plt.show()


# 7.8 Share of succesful projects by goal (equal intervals)
blue_color = sns.color_palette("deep")[0]
cf_cleaned['goal_bin'] = pd.cut(cf_cleaned['c_Goal'], bins=10)
bin_stats = cf_cleaned.groupby('goal_bin').agg(
    success_rate=('success', 'mean'),
    bin_center=('c_Goal', 'mean'),
    count=('success', 'count')).reset_index()
bin_edges = cf_cleaned['goal_bin'].cat.categories
bin_centers = [(b.left + b.right) / 2 for b in bin_edges]
counts = bin_stats['count']
fig, ax1 = plt.subplots(figsize=(10, 6))
ax2 = ax1.twinx()
ax2.bar(bin_centers, counts,
        width=(bin_edges[0].right - bin_edges[0].left)*0.9,
        color='lightgray', label='Počet projektů',
        zorder=1)
ax2.set_ylabel("Počet projektů v intervalu", color='gray')
ax2.tick_params(axis='y', labelcolor='gray')
# Bring ax1 to front
ax1.set_zorder(ax2.get_zorder() + 1)
ax1.patch.set_visible(False)  # Hide the background patch of ax1
# Scatterplot on primary axis
blue_color = sns.color_palette("deep")[0]
ax1.scatter(bin_stats['bin_center'], bin_stats['success_rate'],
            color=blue_color, s=50, label='Úspěšnost', zorder=2)
ax1.set_xlabel("Reálná cílová částka (Kč)")
ax1.set_ylabel("Podíl úspěšných projektů v intervalu", color="black")
ax1.tick_params(axis='y', labelcolor="black")
fig.legend(loc="upper right")
# plt.title("Úspěšnost vs. cílová částka (rovnoměrné intervaly)")
plt.tight_layout()
plt.show()
plt.clf()


# 7.9 Share of succesful projects by goal (equal intervals) by category
# Bin the c_Goal variable into 10 equal-width intervals
cf_cleaned['goal_bin'] = pd.cut(cf_cleaned['c_Goal'], bins=10)
# Calculate success rates and counts per bin and category
bin_stats = cf_cleaned.groupby(['goal_bin', 'Category']).agg(
    success_rate=('success', 'mean'),
    count=('success', 'count')
).reset_index()
# Determine bin centers for plotting
bin_edges = cf_cleaned['goal_bin'].cat.categories
bin_centers = [(interval.left + interval.right) / 2 for interval in bin_edges]
bin_center_mapping = dict(zip(bin_edges, bin_centers))
bin_stats['bin_center'] = bin_stats['goal_bin'].map(bin_center_mapping)
# Calculate total counts per bin for the histogram
bin_counts = cf_cleaned.groupby('goal_bin').size().reset_index(name='count')
bin_counts['bin_center'] = bin_counts['goal_bin'].map(bin_center_mapping)
# Create the plot
fig, ax1 = plt.subplots(figsize=(10, 6))
ax2 = ax1.twinx()
# Plot the histogram on the secondary y-axis
ax2.bar(bin_counts['bin_center'], bin_counts['count'],
        width=(bin_edges[0].right - bin_edges[0].left) * 0.9,
        color='lightgray', label='Počet projektů', zorder=1)
ax2.set_ylabel("Počet projektů v intervalu", color='gray')
ax2.tick_params(axis='y', labelcolor='gray')
# Bring ax1 to the front
ax1.set_zorder(ax2.get_zorder() + 1)
ax1.patch.set_visible(False)  # Hide the background patch of ax1
# Plot the scatter plot on the primary y-axis
sns.scatterplot(data=bin_stats, x='bin_center', y='success_rate',
                hue='Category', style='Category', s=50, ax=ax1, zorder=2)
# Customize 
ax1.set_xlabel("Reálná cílová částka (Kč)")
ax1.set_ylabel("Podíl úspěšných projektů v intervalu", color="black")
ax1.tick_params(axis='y', labelcolor="black")
ax1.legend(title='Kategorie projektu', loc='upper left')
plt.title("Úspěšnost podle cílové částky (rovnoměrné intervaly)")

# Place the legend below the x-axis
ax1.legend(loc='upper center', bbox_to_anchor=(0.5, -0.15), ncol=3)
plt.tight_layout()
plt.show()


# 7.10 Share of succesful projects by interval, size of bubble = number of observation
# Define number of bins
n_bins = 10
# Create equal-width bins for c_Goal
cf_cleaned['goal_bin'] = pd.cut(cf_cleaned['c_Goal'], bins=n_bins)
# Compute success rate in each bin
bin_stats = cf_cleaned.groupby('goal_bin').agg(
    success_rate=('success', 'mean'),
    bin_center=('c_Goal', 'mean'),
    count=('success', 'count')
).reset_index()
# Plotting
plt.figure(figsize=(10, 6))
plt.scatter(bin_stats['bin_center'],
            bin_stats['success_rate'], s=bin_stats['count'], alpha=0.7)
plt.xlabel('Cílová částka (Kč)')
plt.ylabel('Podíl úspěšných projektů')
plt.title('Úspěšnost podle cílové částky (rovnoměrné intervaly)')
plt.grid(True)
plt.tight_layout()
plt.show()


# 7.11 Ratio of succesful projects over Campaign Goal Quantiles by category G1.2bb - ln
# (real value)
# Calculate quantiles
cf_cleaned['c_Goal_quantile'] = pd.qcut(cf_cleaned['ln_c_Goal'],
                                        q=10, labels=False)
# Group by quantiles and calculate average success rate
quantile_success_category = cf_cleaned.groupby(['ln_c_Goal_quantile',
                                               "Category"
                                                ])['success'].mean().reset_index()
# Extract quantile intervals
quantile_intervals = sorted(pd.qcut(cf_cleaned['ln_c_Goal'],
                                    q=10, duplicates='drop').unique())
quantile_intervals_rounded = [pd.Interval(round(interval.left),
                                          round(interval.right)) for interval in quantile_intervals]
# Plot the data
plt.figure(i, figsize=(10, 6))
sns.lineplot(data=quantile_success_category,
             x='ln_c_Goal_quantile', y='success', hue="Category",
             marker='o', linewidth=3)
plt.xlabel('Decil (logaritmu) požadované cílové částky')
plt.ylabel('Podíl úspěšných projektů')
# plt.title('Ratio of succesful projects over Campaign Goal Quantiles by ategory')
# Set xticks to show all quantiles
plt.xticks(ticks=quantile_success_category['ln_c_Goal_quantile'].unique(),
           labels=[str(interval) for interval in quantile_intervals_rounded],
           rotation=45)
# add quantile intervals
# for i, interval in enumerate(quantile_intervals):
#   plt.text(i, quantile_success['success'][i], str(interval), horizontalalignment='center', verticalalignment='bottom')
plt.show()
plt.clf()
i += 1



# 7.12 Ratio of succesful projects over Campaign Goal Quantiles by category G1.2bb - ln
# (nominal value)
# Calculate quantiles
cf_cleaned['Goal_quantile'] = pd.qcut(cf_cleaned['Goal'],
                                      q=10, labels=False,
                                      duplicates="drop")

# Group by quantiles and calculate average success rate
quantile_success_category = cf_cleaned.groupby(['Goal_quantile',
                                               "Category"
                                                ])['success'].mean(
).reset_index()

# Extract quantile intervals
quantile_intervals = sorted(pd.qcut(cf_cleaned['Goal'],
                                    q=10, duplicates='drop').unique())
quantile_intervals_rounded = [pd.Interval(round(interval.left),
                                          round(interval.right)) for interval in quantile_intervals]

# Plot the data
plt.figure(i, figsize=(10, 6))
sns.lineplot(data=quantile_success_category,
             x='Goal_quantile', y='success', hue="Category",
             marker='o', linewidth=3)
plt.xlabel('Campaign Goal Quantile')
plt.ylabel('Ratio of succesful projects')
plt.title('Ratio of succesful projects over Campaign Goal Quantiles by Category (uncompounded values)')

# Set xticks to show all quantiles
plt.xticks(ticks=quantile_success_category['Goal_quantile'].unique(),
           labels=[str(interval) for interval in quantile_intervals_rounded],
           rotation=45)
# add quantile intervals
# for i, interval in enumerate(quantile_intervals):
#   plt.text(i, quantile_success['success'][i], str(interval), horizontalalignment='center', verticalalignment='bottom')

plt.show()
plt.clf()
i += 1

# 7.13 Ratio of succesful projects over Campaign Goal Quantiles by category  !Suspicious, probably includes a mistake
# added quantile sizes
# (real value value)
# could be also done as a heatmap
cf_cleaned['c_Goal_quantile'] = pd.qcut(cf_cleaned['c_Goal'], q=10).astype(str)
# Calculate average success for each quantile and category
avg_success = cf_cleaned.groupby(
    ['c_Goal_quantile', 'Category'])['success'].mean().reset_index()
# Create lineplot with histogram in the background
fig, ax1 = plt.subplots(figsize=(10, 6))
# Plot histogram on secondary y-axis - turn off for thesis:
ax2 = ax1.twinx()
sns.histplot(cf_cleaned['c_Goal_quantile'], bins=10, ax=ax2,
             alpha=0.3)
# ,hue = "Category", multiple = "stack")
ax2.set_ylabel('Number of Observations')
sns.lineplot(data=avg_success, x="c_Goal_quantile", y="success",
             hue="Category", marker="o", linewidth=3, ax=ax1)
ax1.set_xlabel('c_Goal Quantile')
ax1.set_ylabel('Success Rate')
ax1.set_xticklabels(ax1.get_xticklabels(), rotation=45)
plt.show()
plt.clf()
i += 1


# 7.14 Ratio of succesful projects over Campaign Goal Quantiles by category G1.2bb - ln
# (real value, includes outliers)
# let this be token to remind us to use the reasonable cf_cleaned
cf['c_Goal_quantile'] = pd.qcut(cf['c_Goal'], q=10).astype(str)
avg_success = cf.groupby(
    ['c_Goal_quantile', 'Category'])['success'].mean().reset_index()
plt.figure(i)
sns.lineplot(data=avg_success, x="c_Goal_quantile", y="success",
             hue="Category", marker="o", linewidth=3)
plt.xticks(rotation=45)
# plt.xlabel('c')
# plt.ylabel('c_Goal')
plt.title("Based on untreated data including outliers")
plt.show()
plt.clf()
i += 1

# 7.15 Ratio of succesful projects over Campaign Goal by equal intervals by category 
# (real value)
cf_cleaned['c_Goal_bin'] = pd.cut(cf_cleaned['c_Goal'], bins=10)
# Calculate average success for each quantile and category
avg_success = cf_cleaned.groupby(
    ['c_Goal_bin', 'Category'])['success'].mean().reset_index()

# Sort the bins based on the values of the original c_goal variable
avg_success['c_Goal_bin'] = pd.Categorical(avg_success['c_Goal_bin'],
                                           categories=sorted(cf_cleaned['c_Goal_bin'].unique(),
                                           key=lambda x: x.left), ordered=True)
# necessary conversion for plotting:
cf_cleaned['c_Goal_bin'] = cf_cleaned['c_Goal_bin'].astype(str)
avg_success['c_Goal_bin'] = avg_success['c_Goal_bin'].astype(str)
# Create lineplot with histogram in the background
fig, ax1 = plt.subplots(figsize=(10, 6))
# Plot histogram on secondary y-axis - turn off for thesis:
ax2 = ax1.twinx()
sns.histplot(cf_cleaned['c_Goal_bin'],
             ax=ax2, alpha=0.3)
ax2.set_ylabel('Number of Observations')

sns.lineplot(data=avg_success, x="c_Goal_bin", y="success",
             hue="Category", marker="o", linewidth=3, ax=ax1)
ax1.set_xlabel('c_Goal bin')
ax1.set_ylabel('Success Rate')
ax1.set_xticklabels(ax1.get_xticklabels(), rotation=-45)
plt.show()
plt.clf()
i += 1


# 7.16 Bar plot of share of categories in Goal quantiles
plt.figure(i)
sns.histplot(x=cf_cleaned['c_Goal_quantile'],
             hue=cf_cleaned["Category"])
plt.xticks(rotation=45)
plt.show()
plt.clf()
i += 1

#%% Comparison plots


# 7.17 Comparison plot #1 - Goal on succes per category with linear trend
# Prepare the figure
fig, axes = plt.subplots(1, 2, figsize=(14, 6), sharey=True)

# Sort categories alphabetically
category_list = sorted(cf_cleaned['Category'].unique())
palette = sns.color_palette("tab10", n_colors=len(category_list))
category_colors = dict(zip(category_list, palette))

# Plot 1: Linear trends
for cat in category_list:
    subset = cf_cleaned[cf_cleaned['Category'] == cat]
    sns.regplot(
        data=subset, x='c_Goal', y='success',
        ax=axes[0], scatter=True, ci=None,
        label=cat, color=category_colors[cat]
    )
axes[0].set_title("Lineární odhad závislosti")
axes[0].set_xlabel("Cílová částka (Kč v hodnotě k 2024)")
axes[0].set_ylabel("Podíl úspěšných projektů")

# Plot 2: Linear trends on log(goal)
for cat in category_list:
    subset = cf_cleaned[cf_cleaned['Category'] == cat]
    sns.regplot(
        data=subset, x='ln_c_Goal', y='success',
        ax=axes[1], scatter=True, ci=None,
        label=cat, color=category_colors[cat]
    )

axes[1].set_title("Lineární odhad závislosti (log cíl)")
axes[1].set_xlabel("Logaritmus cílové částky (v Kč)")
axes[1].set_ylabel("")

# Shared legend at the bottom center
handles, labels = axes[1].get_legend_handles_labels()
fig.legend(
    handles, labels, title="Kategorie",
    loc='lower center', bbox_to_anchor=(0.5, -0.05), ncol=4
)

# plt.tight_layout()
plt.subplots_adjust(bottom=0.2)  # Make room for legend
plt.show()


# 7.18 Comparison plot #2 - Goal on succes per category with logistic trend
# Prepare the figure
fig, axes = plt.subplots(1, 2, figsize=(14, 6), sharey=True)

# Sort categories alphabetically
category_list = sorted(cf_cleaned['Category'].unique())
palette = sns.color_palette("tab10", n_colors=len(category_list))
category_colors = dict(zip(category_list, palette))

# Plot 2: Logistic trends
for cat in category_list:
    subset = cf_cleaned[cf_cleaned['Category'] == cat]
    sns.regplot(
        data=subset, x='c_Goal', y='success',
        ax=axes[0], scatter=True, ci=None, logistic=True,
        label=cat, color=category_colors[cat]
    )

axes[0].set_title("Logistický odhad závislosti")
axes[0].set_xlabel("Cílová částka (Kč v hodnotě k 2024)")
axes[0].set_ylabel("Podíl úspěšných projektů")

# Plot 2: Logistic trends on log(goal)
for cat in category_list:
    subset = cf_cleaned[cf_cleaned['Category'] == cat]
    sns.regplot(
        data=subset, x='ln_c_Goal', y='success',
        ax=axes[1], scatter=True, ci=None, logistic=True,
        label=cat, color=category_colors[cat]
    )

axes[1].set_title("Logistický odhad závislosti (log cíl)")
axes[1].set_xlabel("Logaritmus cílové částky (v Kč)")
axes[1].set_ylabel("")

# Shared legend at the bottom center
handles, labels = axes[1].get_legend_handles_labels()
fig.legend(
    handles, labels, title="Kategorie",
    loc='lower center', bbox_to_anchor=(0.5, -0.05), ncol=4
)

# plt.tight_layout()
plt.subplots_adjust(bottom=0.2)  # Make room for legend
plt.show()


# 7.19 alternate trend lines for success by category
# logistic estimation
plt.figure(i)
sns.lmplot(data=cf_cleaned, x='c_Goal', y='success',
           col='Category', logistic=True, col_wrap=3, ci=None)
plt.show()
plt.clf()
i += 1


# 7.20 LOESS per category (normal scale)
# Set desired figure dimensions (10:6 ratio = aspect 10/6 ≈ 1.67)
g = sns.lmplot(
    data=cf_cleaned,
    x='c_Goal', y='success',
    hue='Category',
    lowess=True, scatter=True, ci=None,
    height=6, aspect=1.67,
    legend=False  # Turn off default legend
)
# Get handles and labels for the legend from the axes
handles, labels = g.ax.get_legend_handles_labels()
# Add custom legend with title
g.ax.legend(handles, labels, title="Kategorie",
            loc='upper right', bbox_to_anchor=(1, 1))
# Axis labels
g.set_axis_labels("Cílová částka", "Podíl úspěšných projektů")
plt.tight_layout()
plt.show()


# 7.21 LOESS per category (logarithmic scale)
g = sns.lmplot(
    data=cf_cleaned,
    x='ln_c_Goal', y='success',
    hue='Category',
    lowess=True, scatter=True, ci=None,
    height=6, aspect=1.67,
    legend=False  # Turn off default legend
)
# Get handles and labels for the legend from the axes
handles, labels = g.ax.get_legend_handles_labels()
# Add custom legend with title
g.ax.legend(handles, labels, title="Kategorie",
            loc='upper right', bbox_to_anchor=(1, 1))
# Axis labels
g.set_axis_labels("Cílová částka", "Podíl úspěšných projektů")
plt.tight_layout()
plt.show()


# 7.22 Comparison plot - LOESS whole sample (real value)
# Create side-by-side plots
fig, axes = plt.subplots(1, 2, figsize=(12, 6), sharey=True)

# Plot 1: Success vs. Goal (c_Goal)
sns.regplot(
    data=cf_cleaned, x='c_Goal', y='success',
    lowess=True, scatter=True, ci=None, ax=axes[0], scatter_kws={'alpha': 0.4}
)
axes[0].set_xlabel('Cílová částka (v Kč k roku 2024)')
axes[0].set_ylabel('Podíl úspěšných projektů')
axes[0].set_title('LOESS: Cílová částka')

# Plot 2: Success vs. ln(Goal)
sns.regplot(
    data=cf_cleaned, x='ln_c_Goal', y='success',
    lowess=True, scatter=True, ci=None, ax=axes[1], scatter_kws={'alpha': 0.4}
)
axes[1].set_xlabel('Logaritmus reálné cílové částky')
axes[1].set_ylabel('')  # already shared on the left
axes[1].set_title('LOESS: Logaritmus cílové částky')

plt.tight_layout()
plt.show()

# 7.23 Plot of demanded sum for all categories in years
plt.figure(i, figsize=(12, 8))
sns.boxplot(data=cf_cleaned, x='Year', y='Goal', hue='Category')
plt.xlabel('Year')
plt.ylabel('Goal')
plt.title('Flying Box Plots of Goal by Category and Year')
plt.legend(title='Category')
plt.show()
plt.clf()
i += 1


# 7.24 Plot of demanded sum for all categories in years cleaned df
plt.figure(i, figsize=(12, 8))
sns.boxplot(data=cf_cleaned, x='Year', y='Goal', hue='Category')
plt.xlabel('Year')
plt.ylabel('Goal')
plt.title('Flying Box Plots of Goal by Category and Year')
plt.legend(title='Category')
plt.show()
plt.clf()
i += 1

# 7.25 Plot of demanded PV sum for all categories in years cleaned df
plt.figure(i, figsize=(12, 8))
sns.boxplot(data=cf_cleaned, x='Year', y='c_Goal', hue='Category')
plt.xlabel('Year')
plt.ylabel('c_Goal')
plt.title('Flying Box Plots of PV Goal by Category and Year')
plt.legend(title='Category')
plt.show()
plt.clf()
i += 1

# %% 8 Modelling
# 8 Modelling

# Choose which version of the df we work with
df = cf_cleaned
# according to Cook´s distance, no datapoint seems to be
# too influential even in the uncleaned dataset
# according to Mr. Löster, it is better to use the cleaned dataset

# %% 8.1 General model uncompounded
# 8.1 General model uncompounded

formula_g = """success ~ ln_Goal + C(Category, Sum) + n_rewards + N_videos 
 + min_contribution + max_contribution + med_contribution + ln_descr_length
 + Year_number + C(Month, Sum)"""

model_g = smf.logit(formula_g, data=df).fit()
print(model_g.summary())


# Get the summary as a dataframe
summary_df = model_g.summary2().tables[1]

# Export to Excel
file = 'Model_g_summary.xlsx'
# summary_df.to_excel(file)
# print(f"Summary table has been saved to '{file}'.")


# Quality check
# using the cleaned dataset decreases the explanatory power of the model, unexpectedly
mcfadden_r2_g = 1 - (model_g.llf / model_g.llnull)
print(f"McFadden's R2 of general model: {mcfadden_r2_g}")


#%% Something extra to model (1)
# post-check for outliers
influence_g = model_g.get_influence()
df["I_Cooks_d_g"] = influence_g.cooks_distance[0]
df["I_Leverage_g"] = influence_g.hat_matrix_diag

# Calculate threshold
k = model_g.df_model  # Number of predictors (excluding intercept)
n = model_g.nobs  # Number of observations
threshold_g = (2 * (k + 1)) / n

# Identify high leverage points
high_leverage_points = df[df["I_Leverage_g"] > threshold_g]
print("High Leverage Points (general model):")
print(high_leverage_points)

# %% 8.2 General model compounded
# 8.2 General model compounded
formula_compounded = """success ~ ln_c_Goal + C(Category, Sum) + n_rewards + N_videos 
 + c_min_contribution + c_max_contribution + c_med_contribution + ln_descr_length
 + Year_number + C(Month, Sum)"""
# include categories?

f_cols_c = ["ln_c_Goal", "n_rewards", "N_videos", "c_min_contribution",
            "c_max_contribution", "c_med_contribution", "ln_descr_length",
            "Year_number", "Month"]

model_c = smf.logit(formula_compounded, data=df).fit()


print(model_c.summary())

# Quality check
mcfadden_r2_c = 1 - (model_c.llf / model_c.llnull)
print(f"McFadden's R2 of general compounded model: {mcfadden_r2_c}")

#%% Something extra for Model (2)

# post-check for outliers
# influence_c = model_c.get_influence()
# df["I_Cooks_d_c"] = influence_c.cooks_distance[0]
# df["I_Leverage_c"] = influence_c.hat_matrix_diag
# df["I_Stud_resids_c"] = resid_studentized_external

# Calculate threshold
# k = model_c.df_model  # Number of predictors (excluding intercept)
# n = model_c.nobs  # Number of observations
# threshold_c = (2 * (k + 1)) / n

# Identify high leverage points
# leverage_points = df[df["I_Leverage_c"] > threshold_c]
# print("High Leverage Points (general model):")
# print(high_leverage_points)

# Checking randomness of residuals
# Get predicted probabilities
df["predicted"] = model_c.predict()

# Compute deviance residuals
df["deviance_residuals"] = model_c.resid_dev

# Bin predicted probabilities into 10 groups
df["bin"] = pd.qcut(df["predicted"], q=10, duplicates="drop")

# Compute mean predicted probability and mean residual for each bin
binned_data = df.groupby("bin").agg(
    {"predicted": "mean", "deviance_residuals": "mean"}).reset_index()

# Plot binned residuals
plt.scatter(binned_data["predicted"],
            binned_data["deviance_residuals"], color="blue")
plt.axhline(y=0, color="red", linestyle="--")
plt.xlabel("Predicted Probability")
plt.ylabel("Mean Deviance Residuals")
plt.title("Binned Residual Plot")
plt.show()
# the graphic shows possible misspecification or heteroscedasticity

# test for misspecification - linktest
# Logit transformation
df["logit"] = np.log(df["predicted"] / (1 - df["predicted"]))
df["logit_squared"] = df["logit"] ** 2  # Squared term

# Fit logistic regression with logit squared term
linktest_model = smf.logit(
    f"{formula_compounded} + logit_squared", data=df).fit()
print(linktest_model.summary())
# 0 is not in the confidence interval -> misspecification?
# -> probably need to include or omit variable/interaction/whatever

#%% ROC models (1) a (2)
    
# Predict probabilities - ROC data
y_pred_g = model_g.predict(df)
auc_g = roc_auc_score(df['success'], y_pred_g)
fpr_g, tpr_g, _ = roc_curve(df['success'], y_pred_g)

# Predict probabilities - ROC data
y_pred_c = model_c.predict(df)
auc_c = roc_auc_score(df['success'], y_pred_c)
fpr_c, tpr_c, _ = roc_curve(df['success'], y_pred_c)

# ROC Plot Combined (Optional)
plt.figure()
plt.plot(fpr_g, tpr_g, label=f'Základní model (1) (AUC = {
         auc_g:.2f})', color='blue')
plt.plot(fpr_c, tpr_c, label=f'Zúročený model (2) (AUC = {
         auc_c:.2f})', color='orange')

# Diagonal line for random performance
plt.plot([0, 1], [0, 1], 'k--', label='Model na bázi náhody')

# Labels and formatting
# plt.title("ROC Curve Comparison: Základní vs. Compounded Model")
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()


#%% Calculate VIF

df_var_only = df[f_cols_c]

#encode month to effect coding 
# Assuming Month is categorical
df_var_only['Month'] = df_var_only['Month'].astype('category')

# Use patsy to get effect coding
y, X_month = patsy.dmatrices('y ~ C(Month, Sum)', data=df.assign(y=1), return_type='dataframe')
y, X_category = patsy.dmatrices('y ~ C(Category, Sum)', data=df.assign(y=1), return_type='dataframe')

# Drop intercept if needed and the constant 'y'
X_month = X_month.drop(columns=['Intercept'], errors='ignore')
X_category = X_category.drop(columns=['Intercept'], errors='ignore')
# X_coded for future work
X_coded = pd.concat([df_var_only.drop("Month", axis = 1), X_category, X_month], axis = 1)

VIF = calculate_vif(X_coded)
# %% 8.3 Category-wise attempts

# %% 8.3.1 Interactions-including model (compounded)

formula_interactionalist = """success ~ ln_c_Goal + C(Category, Sum) + n_rewards + ln_c_Goal * C(Category, Sum)   + N_videos 
 + c_min_contribution + c_max_contribution + c_med_contribution + ln_descr_length
 + Year_number + C(Month, Sum)"""
# include categories?

# f_cols_i = ["ln_c_Goal", "n_rewards", "N_videos", "c_min_contribution",
#           "c_max_contribution","c_med_contribution", "ln_descr_length",
model_i = smf.logit(formula_interactionalist, data=df).fit()
print(model_i.summary())

# Export model
# export_model_summary_to_word_complete(model_i,"model_i.docs", custom_round = custom_round, eff)
export_model_summary_to_word_auto(
    model=model_i,
    output_path=r"Model_Interaction_Summary2.docx",
    formula=formula_interactionalist,
    data=df,
    custom_round=custom_round
)

export_model_summary_to_word_thesis(
    model=model_i,
    output_path=r"Model_Interaction_Summary3.docx",
    formula=formula_interactionalist,
    data=df,
    custom_round=custom_round)

export_model_summary_to_word_thesis_m(
    model=model_i,
    output_path=r"Model_Interaction_Summary4.docx",
    formula=formula_interactionalist,
    data=df,
    custom_round=custom_round)

# %%

# Quality check
mcfadden_r2_i = 1 - (model_i.llf / model_i.llnull)
print(f"McFadden's R2 of model with interactions: {mcfadden_r2_c}")

# Predict probabilities - ROC data
y_pred_i = model_i.predict(df)
auc_i = roc_auc_score(df['success'], y_pred_i)
fpr_i, tpr_i, _ = roc_curve(df['success'], y_pred_i)

# ROC Plot (Optional)
plt.figure()
plt.plot(fpr_i, tpr_i, label=f'Interaction Model (AUC = {auc_i:.2f})')
plt.plot([0, 1], [0, 1], 'k--')
# plt.title("ROC Curve: Interaction Model")
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()

# Checking randomness of residuals
# Get predicted probabilities
df["predicted"] = model_i.predict()

# Compute deviance residuals
df["deviance_residuals"] = model_i.resid_dev

# Bin predicted probabilities into 10 groups
df["bin"] = pd.qcut(df["predicted"], q=10, duplicates="drop")

# Compute mean predicted probability and mean residual for each bin
binned_data = df.groupby("bin").agg(
    {"predicted": "mean", "deviance_residuals": "mean"}).reset_index()

# Plot binned residuals
plt.figure(i)
plt.scatter(binned_data["predicted"],
            binned_data["deviance_residuals"], color="blue")
plt.axhline(y=0, color="red", linestyle="--")
plt.xlabel("Predicted Probability")
plt.ylabel("Mean Deviance Residuals")
plt.title("Binned Residual Plot")
plt.show()
plt.clf()
i += 1
# the graphic shows possible misspecification or heteroscedasticity

# test for misspecification - linktest
# Logit transformation
df["logit"] = np.log(df["predicted"] / (1 - df["predicted"]))
df["logit_squared"] = df["logit"] ** 2  # Squared term

# Fit logistic regression with logit squared term
linktest_model = smf.logit(
    f"{formula_interactionalist} + logit_squared", data=df).fit()
print(linktest_model.summary())
# 0 is not in the confidence interval -> misspecification?
# -> probably need to include or omit variable/interaction/whatever

# %% Interaction model - curves prediction

# Define the range of ln_c_Goal values
ln_c_goal_range = np.linspace(
    df['ln_c_Goal'].min(), df['ln_c_Goal'].max(), 100)

# Get unique categories
categories = df['Category'].unique()

# Create a DataFrame for prediction
predict_df = pd.DataFrame([
    {'ln_c_Goal': ln_c_goal, 'Category': category}
    for category in categories
    for ln_c_goal in ln_c_goal_range
])

# Set other covariates to their mean or reference values
mean_values = df.mean(numeric_only=True)
for covariate in ['n_rewards', 'N_videos', 'c_min_contribution', 'c_max_contribution',
                  'c_med_contribution', 'ln_descr_length', 'Year_number']:
    predict_df[covariate] = mean_values[covariate]

# For 'Month', set to the most frequent category
predict_df['Month'] = df['Month'].mode()[0]

# Include original categories
predict_df['c_Goal'] = np.exp(predict_df['ln_c_Goal'])


# Predict probabilities
predict_df['predicted_success'] = model_i.predict(predict_df)

# Build plot
# Set the style
sns.set(style="whitegrid")

# Initialize the figure and axes
fig, axes = plt.subplots(1, 2, figsize=(14, 6), sharey=True)

# Define the color palette
category_list = sorted(cf_cleaned['Category'].unique())
palette = sns.color_palette("tab10", n_colors=len(category_list))
category_colors = dict(zip(category_list, palette))

# Plot for original scale (c_Goal)
for category in category_list:
    pred_subset = predict_df[predict_df['Category'] == category]
    axes[0].plot(pred_subset['c_Goal'], pred_subset['predicted_success'],
                 color=category_colors[category], linewidth=2, label=category)  # Plot line first
    subset = df[df['Category'] == category]
    axes[0].scatter(subset['c_Goal'], subset['success'],
                    alpha=0.3, color=category_colors[category])  # Then scatter

axes[0].set_xlabel("Cílová částka (Kč v hodnotě k 2024)")
axes[0].set_ylabel("Odhadovaná pravděpodobnost úspěchu")
axes[0].set_title("Predikce úspěšnosti podle cílové částky")

# Plot for logarithmic scale (ln_c_Goal)
for category in category_list:
    pred_subset = predict_df[predict_df['Category'] == category]
    axes[1].plot(pred_subset['ln_c_Goal'], pred_subset['predicted_success'],
                 color=category_colors[category], linewidth=2, label=category)  # Plot line first
    subset = df[df['Category'] == category]
    axes[1].scatter(subset['ln_c_Goal'], subset['success'],
                    alpha=0.3, color=category_colors[category])  # Then scatter

axes[1].set_xlabel("Logaritmus reálné cílové částky (Kč)")
axes[1].set_title("Predikce úspěšnosti podle logaritmu cílové částky")

# Legend
handles, labels = axes[1].get_legend_handles_labels()
fig.legend(handles, labels, title="Kategorie",
           loc='lower center', bbox_to_anchor=(0.5, -0.05), ncol=4)

plt.subplots_adjust(bottom=0.22)
plt.show()

#%% GAM models data

# Step 1: Copy the data and prepare predictors
df_model = df.copy()

# Encode categorical variables numerically
df_model['Category_code'] = df_model['Category'].cat.codes
df_model['Month_code'] = df_model['Month'].cat.codes

category_map = dict(enumerate(df_model['Category'].cat.categories))
month_map = dict(enumerate(df_model['Month'].cat.categories))


#%%
    
def twinsubplots(model):
    # Create the twin subplots
    fig, axes = plt.subplots(ncols=2, figsize=(12, 5), sharey=True)
    
    # Create range of ln(c_Goal) and convert to raw c_Goal
    ln_goal_range = np.linspace(
        df_model['ln_c_Goal'].min(), df_model['ln_c_Goal'].max(), 200)
    goal_range = np.exp(ln_goal_range)  # back to original c_Goal values

    # Finding the code for December
    december_code = [code for code, name in month_map.items() if name == 12][0]
    
    # Define the color palette
    category_list = sorted(cf_cleaned['Category'].unique())
    palette = sns.color_palette("tab10", n_colors=len(category_list))
    category_colors = dict(zip(category_list, palette))

    # Plot background scatterplots
    for cat in category_list:
        cat_data = df[df['Category'] == cat]
        color = category_colors[cat]
        
        # Left subplot (linear x-axis)
        axes[0].scatter(
            cat_data['c_Goal'], cat_data['success'],
            label=None, color=color, alpha=0.3, s=10)
        
        # Right subplot (log x-axis)
        axes[1].scatter(
            cat_data['c_Goal'], cat_data['success'],
            label=None, color=color, alpha=0.3, s=10)

    # Overlay model predictions
    for cat_code, cat_name in category_map.items():
        # Construct prediction input
        X_pred = np.zeros((len(ln_goal_range), X.shape[1]))
        X_pred[:, 0] = ln_goal_range
        X_pred[:, 1] = cat_code
        X_pred[:, 2:] = 0
        X_pred[:, 9] = december_code

        pred_probs = model.predict_proba(X_pred)
        color = category_colors[cat_name]

        # Plot on left (linear x-axis)
        axes[0].plot(goal_range, pred_probs, label=cat_name, color=color)

        # Plot on right (logarithmic x-axis)
        axes[1].plot(goal_range, pred_probs, label=cat_name, color=color)

    # Left subplot (linear scale)
    axes[0].set_xlabel("Cílová částka (Kč v hodnotě k 2024)")
    axes[0].set_ylabel("Odhadovaná pravděpodobnost úspěchu")
    axes[0].set_title("Předikce úspěšnosti podle cílové částky")
    axes[0].grid(True)

    # Right subplot (log scale)
    axes[1].set_xscale('log')
    axes[1].set_xlabel("Logaritmus cílové částky (v Kč)")
    axes[1].set_title("Predikce úspěšnosti podle logaritmu cílové částky")
    axes[1].grid(True)

    # Remove legends from individual axes if they exist
    if axes[0].legend_: axes[0].legend_.remove()
    if axes[1].legend_: axes[1].legend_.remove()

    # Shared legend below both plots
    handles, labels = axes[1].get_legend_handles_labels()
    fig.legend(
        handles, labels, title="Kategorie",
        loc='lower center', bbox_to_anchor=(0.5, -0.05), ncol=4
    )

    plt.subplots_adjust(bottom=0.3)
    plt.show()



# %% 8.3.2 The GAM (tensor smoothed)
#%% Tensor GAM

# Step 2: Select predictor variables
X = df_model[[
    'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
    'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
    'ln_descr_length', 'Year_number', 'Month_code'
]].values

y = df_model['success'].values

# Optional: Standardize continuous variables (recommended)
# stáhne střední hodnotu na nula a std. na 1 - standardizuje rozdělení (a to i nenormální rozdělení)
scaler = StandardScaler()
X[:, 2:9] = scaler.fit_transform(X[:, 2:9])
# It is better if the scaled variables are closer to norm.dist

# We scaled all covariables, as we are not going to extract
# their values out of the model anyway, and this will
# a) make them numerically stable for fitting - penalization
# will behave better if they have similar range
# b) it allows for interpretting the model as showing effects

# Step 3: Fit the GAM
gam_te_free = LogisticGAM(
    te(0, 1, lam=[0,0]) +      # tensor smooth: ln_c_Goal × Category_code
    f(1) +          # Category baseline effect
    l(2) +          # n_rewards
    l(3) +          # N_videos
    l(4) +          # c_min_contribution
    l(5) +          # c_max_contribution
    l(6) +          # c_med_contribution
    l(7) +          # ln_descr_length
    l(8) +          # Year_number
    f(9)            # Month (as categorical effect)
    , fit_intercept=True).fit(X, y)

#gam_te.gridsearch(X, y) #gridsearch vyjde vždy zhruba stejně nehledě na mód

print("Model fitted successfully.")
print(gam_te_free.summary())

# Log-likelihood of GAM (already fit)
loglik_gam_te_free = gam_te_free.loglikelihood(X, y)

# Null model: only intercept
X_null = np.ones((X.shape[0], 1))  # column of ones
gam_null = LogisticGAM().fit(X_null, y)
loglik_null = gam_null.loglikelihood(X_null, y)

# Checking model quality
# Compute McFadden R²
mcfadden_r2_gam_te_free = 1 - (loglik_gam_te_free / loglik_null)
print(f"McFadden's R2 for GAM_te_free: {mcfadden_r2_gam_te_free:.4f}")

# Predict probabilities for the full dataset
y_pred_gam_te_free = gam_te_free.predict_proba(X)

# Compute AUC
auc_gam_te_free = roc_auc_score(y, y_pred_gam_te_free)
print(f"AUC for GAM_te_free: {auc_gam_te_free:.4f}")

# Optional: Plot the ROC curve
fpr, tpr, _ = roc_curve(y, y_pred_gam_te_free)

plt.figure()
plt.plot(fpr, tpr, color='blue', lw=2,
         label=f'ROC curve (AUC = {auc_gam_te_free:.2f})')
plt.plot([0, 1], [0, 1], color='gray', lw=2, linestyle='--')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('GAM: Receiver Operating Characteristic (ROC) Curve')
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()


twinsubplots(gam_te_free)

#%% GAM tensor #2 (Optimized lambda)

from itertools import product

# Step 1: Copy the data and prepare predictors
df_model = df.copy()

# Encode categorical variables numerically
df_model['Category_code'] = df_model['Category'].cat.codes
df_model['Month_code'] = df_model['Month'].cat.codes

category_map = dict(enumerate(df_model['Category'].cat.categories))
month_map = dict(enumerate(df_model['Month'].cat.categories))

# Step 2: Select predictor variables
X = df_model[[
    'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
    'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
    'ln_descr_length', 'Year_number', 'Month_code'
]].values

y = df_model['success'].values

# Standardize continuous variables
scaler = StandardScaler()
X[:, 2:9] = scaler.fit_transform(X[:, 2:9])

# Step 3: Grid search for best lambda
lam_space = np.logspace(-3, 2, 6)  # [0.001, 0.01, 0.1, 1, 10, 100]
best_score = -np.inf
best_model = None
best_lam = None

for lam1, lam2 in product(lam_space, repeat=2):
    gam_te_try = LogisticGAM(
        te(0, 1, lam=[lam1, lam2]) +
        f(1) +
        l(2) + l(3) + l(4) + l(5) + l(6) + l(7) + l(8) +
        f(9), fit_intercept=True
    )
    gam_te_try.fit(X, y)
    score = gam_te_try.loglikelihood(X, y)
    if score > best_score:
        best_score = score
        best_model = gam_te_try
        best_lam = [lam1, lam2]

gam_te = best_model
print(f"Model fitted successfully with lam = {best_lam}")

# Log-likelihood and McFadden R²
loglik_gam_te = gam_te.loglikelihood(X, y)
X_null = np.ones((X.shape[0], 1))
gam_null = LogisticGAM().fit(X_null, y)
loglik_null = gam_null.loglikelihood(X_null, y)
mcfadden_r2_gam_te = 1 - (loglik_gam_te / loglik_null)
print(f"McFadden's R2 for GAM_te: {mcfadden_r2_gam_te:.4f}")

# Predict probabilities and compute AUC
y_pred_gam_te = gam_te.predict_proba(X)
auc_gam_te = roc_auc_score(y, y_pred_gam_te)
print(f"AUC for GAM_te: {auc_gam_te:.4f}")

# Create the twin subplots
fig, axes = plt.subplots(ncols=2, figsize=(12, 5), sharey=True)

# Create range of ln(c_Goal) and convert to raw c_Goal
ln_goal_range = np.linspace(
    df_model['ln_c_Goal'].min(), df_model['ln_c_Goal'].max(), 200)
goal_range = np.exp(ln_goal_range)  # back to original c_Goal values

# Finding the code for december
december_code = [code for code, name in month_map.items() if name == 12][0]

for cat_code, cat_name in category_map.items():
    # Construct prediction input
    X_pred = np.zeros((len(ln_goal_range), X.shape[1]))
    X_pred[:, 0] = ln_goal_range
    X_pred[:, 1] = cat_code
    X_pred[:, 2:] = 0
    X_pred[:, 9] = december_code

    pred_probs = gam_te.predict_proba(X_pred)

    # Plot on left (linear x-axis)
    axes[0].plot(goal_range, pred_probs, label=cat_name)

    # Plot on right (logarithmic x-axis)
    axes[1].plot(goal_range, pred_probs, label=cat_name)

# Left subplot (linear scale)
axes[0].set_xlabel("Cílová částka (Kč v hodnotě k 2024)")
axes[0].set_ylabel("Odhadovaná pravděpodobnost úspěchu")
axes[0].set_title("Lineární měřítko")
axes[0].grid(True)

# Right subplot (log scale)
axes[1].set_xscale('log')
axes[1].set_xlabel("Logaritmus cílové částky (v Kč)")
axes[1].set_title("Logaritmické měřítko")
axes[1].grid(True)
axes[1].legend()

# Remove legends from both axes
axes[0].legend_.remove() if axes[0].legend_ else None
axes[1].legend_.remove() if axes[1].legend_ else None

# Shared legend below both plots
handles, labels = axes[1].get_legend_handles_labels()
fig.legend(
    handles, labels, title="Kategorie",
    loc='lower center', bbox_to_anchor=(0.5, -0.05), ncol=4
)

# Adjust spacing to make room for legend
plt.subplots_adjust(bottom=0.3)
# plt.tight_layout()
plt.show()


#%% GAM tensor #3 (Rigorously optimized lambda)

from pygam import LogisticGAM, s, f, l, te
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import KFold
import numpy as np

# Step 1: Copy the data and prepare predictors
df_model = df.copy()

# Encode categorical variables numerically
df_model['Category_code'] = df_model['Category'].cat.codes
df_model['Month_code'] = df_model['Month'].cat.codes

category_map = dict(enumerate(df_model['Category'].cat.categories))
month_map = dict(enumerate(df_model['Month'].cat.categories))

# Step 2: Select predictor variables
X = df_model[[ 
    'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
    'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
    'ln_descr_length', 'Year_number', 'Month_code'
]].values

y = df_model['success'].values

# Standardize continuous covariates (recommended)
scaler = StandardScaler()
X[:, 2:9] = scaler.fit_transform(X[:, 2:9])

# Step 3: Define the model structure
gam_te3 = LogisticGAM(
    te(0, 1) +       # tensor smooth: ln_c_Goal × Category_code
    f(1) +           # Category baseline effect
    l(2) +           # n_rewards
    l(3) +           # N_videos
    l(4) +           # c_min_contribution
    l(5) +           # c_max_contribution
    l(6) +           # c_med_contribution
    l(7) +           # ln_descr_length
    l(8) +           # Year_number
    f(9),            # Month as factor
    fit_intercept=True
)

# Step 4: Run 5-fold CV gridsearch for λ
gam_te3.gridsearch(X, y, objective= "AIC")

# Step 5: Evaluate and report lambda and model diagnostics
print("Model fitted successfully using 5-fold CV for λ selection.")
print(f"Selected lambda(s): {gam_te3.lam}")

# Compute AIC and BIC
loglik = gam_te3.loglikelihood(X, y)
edf = gam_te3.statistics_['edof']
n = X.shape[0]
aic = 2 * edf - 2 * loglik
bic = np.log(n) * edf - 2 * loglik

print(f"Effective Degrees of Freedom (EDF): {edf:.2f}")
print(f"AIC: {aic:.2f}")
print(f"BIC: {bic:.2f}")

# Optional: Print detailed model summary
print(gam_te3.summary())

# Checking model quality
# Log-likelihood of GAM (already fit)
loglik_gam_te3 = gam_te3.loglikelihood(X, y)
# Compute McFadden R²
mcfadden_r2_gam_te3 = 1 - (loglik_gam_te3 / loglik_null)
print(f"McFadden's R2 for GAM: {mcfadden_r2_gam_te3:.4f}")

# Predict probabilities for the full dataset
y_pred_gam_te3 = gam_te3.predict_proba(X)

# Compute AUC
auc_gam = roc_auc_score(y, y_pred_gam_te3)
print(f"AUC for GAM: {auc_gam:.4f}")

# Optional: Plot the ROC curve
fpr, tpr, _ = roc_curve(y, y_pred_gam_te3)

plt.figure()
plt.plot(fpr, tpr, color='blue', lw=2,
         label=f'ROC curve (AUC = {auc_gam:.2f})')
plt.plot([0, 1], [0, 1], color='gray', lw=2, linestyle='--')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('GAM: Receiver Operating Characteristic (ROC) Curve')
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()

twinsubplots(gam_te3)

#%% GAM tensor custom lambda

from pygam import LogisticGAM, s, f, l, te
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import KFold
import numpy as np

# Step 1: Copy the data and prepare predictors
df_model = df.copy()

# Encode categorical variables numerically
df_model['Category_code'] = df_model['Category'].cat.codes
df_model['Month_code'] = df_model['Month'].cat.codes

category_map = dict(enumerate(df_model['Category'].cat.categories))
month_map = dict(enumerate(df_model['Month'].cat.categories))

# Step 2: Select predictor variables
X = df_model[[ 
    'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
    'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
    'ln_descr_length', 'Year_number', 'Month_code'
]].values

y = df_model['success'].values

# Standardize continuous covariates (recommended)
scaler = StandardScaler()
X[:, 2:9] = scaler.fit_transform(X[:, 2:9])

# Step 3: Define the model structure
gam_te4 = LogisticGAM(
    te(0, 1, lam = [0.1, 0.1]) +       # tensor smooth: ln_c_Goal × Category_code
    f(1) +           # Category baseline effect
    l(2) +           # n_rewards
    l(3) +           # N_videos
    l(4) +           # c_min_contribution
    l(5) +           # c_max_contribution
    l(6) +           # c_med_contribution
    l(7) +           # ln_descr_length
    l(8) +           # Year_number
    f(9),            # Month as factor
    fit_intercept=True
).fit(X, y)

# Step 4: Run 5-fold CV gridsearch for λ
#gam_te4.gridsearch(X, y)

# Optional: Print detailed model summary
print(gam_te4.summary())

# Checking model quality
# Log-likelihood of GAM (already fit)
loglik_gam_te4 = gam_te4.loglikelihood(X, y)
# Compute McFadden R²
mcfadden_r2_gam_te4 = 1 - (loglik_gam_te4 / loglik_null)
print(f"McFadden's R2 for GAM: {mcfadden_r2_gam_te4:.4f}")

# Predict probabilities for the full dataset
y_pred_gam_te4 = gam_te4.predict_proba(X)

# Compute AUC
auc_gam_te4 = roc_auc_score(y, y_pred_gam_te4)
print(f"AUC for GAM: {auc_gam_te4:.4f}")

# Optional: Plot the ROC curve
fpr, tpr, _ = roc_curve(y, y_pred_gam_te4)

plt.figure()
plt.plot(fpr, tpr, color='blue', lw=2,
         label=f'ROC curve (AUC = {auc_gam:.2f})')
plt.plot([0, 1], [0, 1], color='gray', lw=2, linestyle='--')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('GAM: Receiver Operating Characteristic (ROC) Curve')
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()


# Create the twin subplots
twinsubplots(gam_te4)

#%% Selecting best lambda gam to equal the logit
# Step 1: Copy the data and prepare predictors
df_model = df.copy()

# Encode categorical variables numerically
df_model['Category_code'] = df_model['Category'].cat.codes
df_model['Month_code'] = df_model['Month'].cat.codes

category_map = dict(enumerate(df_model['Category'].cat.categories))
month_map = dict(enumerate(df_model['Month'].cat.categories))

# Step 2: Select predictor variables
X = df_model[[ 
    'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
    'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
    'ln_descr_length', 'Year_number', 'Month_code'
]].values

y = df_model['success'].values

# Standardize continuous covariates (recommended)
scaler = StandardScaler()
X[:, 2:9] = scaler.fit_transform(X[:, 2:9])

def select_best_lambda_via_r2(X, y, df_model, category_map, month_map,
                              benchmark_R2,
                              lambda_range=np.logspace(-3, 2, 200)):
    r2_benchmark = benchmark_R2
    print(f"Benchmark McFadden R²: {r2_benchmark:.4f}")

    best_lam = None
    best_model = None

    for lam in lambda_range:
        gam = LogisticGAM(
            te(0, 1, lam=lam) +
            f(1) + l(2) + l(3) + l(4) + l(5) + l(6) + l(7) + l(8) + f(9),
            fit_intercept=True
        ).fit(X, y)

        ll_model = gam.loglikelihood(X, y)
        null_gam = LogisticGAM().fit(np.zeros((len(y), 1)), y)
        ll_gam_null = null_gam.loglikelihood(np.zeros((len(y), 1)), y) # Null log-likelihood from deviance
        r2_gam = 1 - (ll_model / ll_gam_null)

        if r2_gam >= r2_benchmark:
            best_lam = lam
            best_model = gam
        else:
            break  # stop if it dips below the benchmark

    if best_model is not None:
        print(f"Selected λ = {best_lam:.4f} with McFadden R² = {r2_gam:.4f}")
    else:
        print("No GAM met or exceeded the benchmark R². Consider extending the lambda range or loosening the criterion.")

    return best_model, best_lam

gam_4c, best_lam = select_best_lambda_via_r2(X, y, df_model, category_map, month_map, benchmark_R2=mcfadden_r2_i)

twinsubplots(gam_4c)

# %% 8.3.2 The Gam - Smoothed

# Step 1: Copy and prepare data again
df_model = df.copy()
df_model['Category_code'] = df_model['Category'].cat.codes
df_model['Month_code'] = df_model['Month'].cat.codes

# Step 2: Create design matrix
X_s = df_model[[  # same variable order as before
    'ln_c_Goal', 'Category_code', 'n_rewards', 'N_videos',
    'c_min_contribution', 'c_max_contribution', 'c_med_contribution',
    'ln_descr_length', 'Year_number', 'Month_code'
]].values
y_s = df_model['success'].values

# Standardize control variables (cols 2 to 8)
scaler_s = StandardScaler()
X_s[:, 2:9] = scaler_s.fit_transform(X_s[:, 2:9])

# Step 3: Fit the GAM with smoothed control variables
gam_s = LogisticGAM(
    te(0, 1, lam = best_lam) +      # smooth interaction: ln_c_Goal × Category_code
    f(1) +          # Category baseline
    s(2) +          # Smoothed control terms
    s(3) +
    s(4) +
    s(5) +
    s(6) +
    s(7) +
    s(8) +
    f(9)            # Month as factor
    , fit_intercept=True).fit(X_s, y_s)

print("Smoothed GAM fitted successfully.")

# Step 4: Evaluate model
loglik_s = gam_s.loglikelihood(X_s, y_s)

# Null model: intercept only
X_null_s = np.ones((X_s.shape[0], 1))
gam_null_s = LogisticGAM().fit(X_null_s, y_s)
loglik_null_s = gam_null_s.loglikelihood(X_null_s, y_s)

# McFadden R²
mcfadden_r2_s = 1 - (loglik_s / loglik_null_s)
print(f"McFadden's R2 for GAM smoothed: {mcfadden_r2_s:.4f}")

# AUC and ROC
y_pred_s = gam_s.predict_proba(X_s)
auc_s = roc_auc_score(y_s, y_pred_s)
print(f"AUC for GAM smoothed: {auc_s:.4f}")

# ROC curve plot
fpr_s, tpr_s, _ = roc_curve(y_s, y_pred_s)

plt.figure()
plt.plot(fpr_s, tpr_s, color='darkorange', lw=2,
         label=f'GAM smoothed (AUC = {auc_s:.2f})')
plt.plot([0, 1], [0, 1], color='gray', linestyle='--')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('GAM Smoothed: ROC Curve')
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()


# Create the twin subplots
twinsubplots(gam_s)

#%% videos plot
import matplotlib.pyplot as plt
import numpy as np

# Get mean values for numeric predictors (standardized ones: columns 2 to 8)
mean_values = X_s[:, 2:9].mean(axis=0)

# Identify code for "Umění a multimédia" and month 12
cat_code = df_model.loc[df_model['Category'] == 'Umění a multimédia', 'Category_code'].iloc[0]
month_code = df_model.loc[df_model['Month'] == 12, 'Month_code'].iloc[0]

# Create matrix for prediction
n_videos_vals = np.array([0, 1, 2, 3])
pred_X = np.zeros((len(n_videos_vals), X_s.shape[1]))

# ln(c_Goal): use mean
pred_X[:, 0] = df_model['ln_c_Goal'].mean()
# Category
pred_X[:, 1] = cat_code
# Control vars (mean values)
pred_X[:, 2:3] = mean_values[0]  # n_rewards
pred_X[:, 4:5] = mean_values[2]  # c_min_contribution
pred_X[:, 5:6] = mean_values[3]  # c_max_contribution
pred_X[:, 6:7] = mean_values[4]  # c_med_contribution
pred_X[:, 7:8] = mean_values[5]  # ln_descr_length
pred_X[:, 8:9] = mean_values[6]  # Year_number
pred_X[:, 9] = month_code

# Now set standardized n_videos
# Standardize like the training data
mean_nv, std_nv = scaler_s.mean_[1], scaler_s.scale_[1]
pred_X[:, 3] = (n_videos_vals - mean_nv) / std_nv

# Predict
probs = gam_s.predict_proba(pred_X)

# Plot
plt.figure(figsize=(8, 5))
bars = plt.bar(n_videos_vals, probs, color=blue_color)

# Add numerical labels above bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width() / 2, height + 0.02,
             f"{height:.2f}", ha='center', va='bottom', fontsize=10)

# Axis settings
plt.xlabel("Počet videí na stránce projektu")
plt.ylabel("Predikovaná pravděpodobnost úspěchu")
plt.title("GAM model vyhlazený: Pravděpodobnost úspěchu podle počtu videí\n(kategorie = Umění a multimédia, měsíc = prosinec)")
plt.xticks(n_videos_vals)        # Show only whole numbers on X-axis
plt.yticks(np.linspace(0, 1, 11))  # Show 0.0 to 1.0 in steps of 0.1
plt.ylim(0, 1)
plt.grid(True, axis='y', linestyle='--', alpha=0.7)
plt.tight_layout()
plt.show()

# %%
# The paired plots final

for cat_code, cat_name in category_map.items():
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(
        8, 6), sharex=True, height_ratios=[2, 1])

    # Step 1: Generate prediction inputs
    ln_goal_range = np.linspace(
        df_model['ln_c_Goal'].min(), df_model['ln_c_Goal'].max())
    X_pred = np.zeros((len(ln_goal_range), X.shape[1]))
    X_pred[:, 0] = ln_goal_range
    X_pred[:, 1] = cat_code
    X_pred[:, 2:] = 0

    # Step 2: Predict probabilities and compute derivative
    probs = gam_s.predict_proba(X_pred)
    derivs = np.gradient(probs, ln_goal_range)

    # Convert X-axis to real goal values for interpretability
    goal_values = np.exp(ln_goal_range)

    # Plot 1: Predicted success probability
    ax1.plot(goal_values, probs, color='blue', lw=2)
    ax1.set_ylabel("Pravděpodobnost úspěchu")
    ax1.set_title(f"{cat_name}: Predikovaná pravděpodobnosts úspěchu v závislosti na cílové částce")
    ax1.grid(True)

    # Plot 2: Derivative of success probability
    ax2.plot(goal_values, derivs, color='green', lw=2)
    ax2.axhline(0, color='gray', linestyle='--')
    ax2.set_xlabel("Cílová částka (v Kč v hodnotě k roku 2024)")  # (log scale)")
    ax2.set_ylabel("Sklon")
    # Log Goal Unit")
    ax2.set_title(
        f"{cat_name}: Změna pravděpodobnosti úspěchu za jednotku cílové částky")
    ax2.grid(True)
    # ax2.set_xscale('log')   # nechat logaritmickou verzi, nebo raději lineární?

    # ticks = [5000, 10000, 20000, 40000, 60000, 100000, 200000, 400000, 800000, 1000000, 1200000, 1400000]
    # ax2.xaxis.set_major_locator(FixedLocator(ticks))
    # shows plain numbers, not scientific notation
    ax2.xaxis.set_major_formatter(ScalarFormatter())
    ax.tick_params(axis='x', rotation=45)  # rotate for readability

    plt.tight_layout()
    plt.show()
    plt.clf()
#%% ROC semiparametric (4c) vs smoothed (5)

# Predict probabilities - ROC data
y_pred_4c = gam_4c.predict_proba(X_s)
auc_4c = roc_auc_score(y_s, y_pred_4c)
fpr_4c, tpr_4c, _ = roc_curve(df['success'], y_pred_4c)

# AUC and ROC
y_pred_s = gam_s.predict_proba(X_s)
auc_s = roc_auc_score(y_s, y_pred_s)
print(f"AUC for GAM smoothed: {auc_s:.4f}")
# ROC curve plot
fpr_s, tpr_s, _ = roc_curve(y_s, y_pred_s)

# ROC Plot Combined (Optional)
plt.figure()
plt.plot(fpr_4c, tpr_4c, label=f'Semiparametrický GAM (4c) (AUC = {
         auc_4c:.2f})', color='blue')
plt.plot(fpr_s, tpr_s, label=f'Neparametrický GAM (5) (AUC = {
         auc_s:.2f})', color='orange')

# Diagonal line for random performance
plt.plot([0, 1], [0, 1], 'k--', label='Model na bázi náhody')

# Labels and formatting
# plt.title("ROC Curve Comparison: Základní vs. Compounded Model")
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.legend(loc='lower right')
plt.grid(True)
plt.tight_layout()
plt.show()

#%% month plot
import matplotlib.pyplot as plt
import pandas as pd

# Define data
data = {
    "Měsíc": ["Leden", "Únor", "Březen", "Duben", "Květen", "Červen", "Červenec",
              "Srpen", "Září", "Říjen", "Listopad", "Prosinec"],
    "Odhad": [-0.13, -0.021, -0.022, 0.031, -0.024, 0.021, -0.104,
              -0.288, -0.213, 0.229, 0.379, 0.142],
    "CI_Low": [-0.366, -0.321, -0.242, -0.184, -0.216, -0.174, -0.306,
               -0.529, -0.461, 0.022, 0.177, None],
    "CI_High": [0.105, 0.278, 0.198, 0.246, 0.169, 0.216, 0.098,
                -0.048, 0.036, 0.437, 0.581, None]
}

df = pd.DataFrame(data)
df["Měsíc"] = pd.Categorical(df["Měsíc"], categories=data["Měsíc"], ordered=True)

# Plot
plt.figure(figsize=(10, 5))
plt.errorbar(
    df["Měsíc"][:-1], df["Odhad"][:-1],
    yerr=[df["Odhad"][:-1] - df["CI_Low"][:-1], df["CI_High"][:-1] - df["Odhad"][:-1]],
    fmt='o', capsize=5, color='tab:blue', markersize=8, linewidth=2
)
plt.axhline(0, color='gray', linestyle='--', linewidth=1)
plt.title("Odhad kategoriálních efektů měsíců s intervaly spolehlivosti (95 %)")
plt.ylabel("Odhad vlivu (logit pravděpodobnosti úspěchu)")
plt.xlabel("Měsíc")
plt.xticks(rotation=45)
plt.tight_layout()
plt.show()

# %% Results

# Model quality assesment
# Compile the metrics into a list of dictionaries
model_summary = [
    {
        "Model": "General Logit",
        "McFadden R2": round(mcfadden_r2_g, 4),
        "AUC": round(auc_gen, 4)
    },
    {
        "Model": "Compounded Logit",
        "McFadden R2": round(mcfadden_r2_c, 4),
        "AUC": round(auc_c, 4)
    },
    {
        "Model": "Interaction Logit (Goal × Category)",
        "McFadden R2": round(mcfadden_r2_i, 4),
        "AUC": round(auc_i, 4)
    },
    {
        "Model": "GAM (Tensor Smooth: Goal × Category)",
        "McFadden R2": round(mcfadden_r2_gam, 4),
        "AUC": round(auc_gam, 4)
    },
    {
        "Model": "Smoothed GAM (Smoothed Control Variables)",
        "McFadden R2": round(mcfadden_r2_s, 4),
        "AUC": round(auc_s, 4)
    }
]

# Create a DataFrame
df_model_summary = pd.DataFrame(model_summary)

# Display the summary table
print(df_model_summary)


# export all the models
for m in [[model_gen, "M_1_Summary.docx", formula_general, "Model (1): Logit v nominálních hodnotách"],
          [model_c, "M_2_Summary.docx", formula_compounded,
              "Model (2): Logit v reálných hodnotách"],
          [model_i, "M_3_Summary.docx", formula_general, "Model (3): Logit s interakcemi"]]:
    export_model_summary_to_word_thesis_m(
        model=m[0],
        output_path=m[1],
        formula=m[2],
        data=df,
        custom_round=custom_round,
        model_name=m[3])


# export model summaries to excel
# export_model_summaries_to_excel()

# %%
# X - Ending

# at he end, we save dataframe as a csv
cf.to_csv("cf.csv", index=False)
